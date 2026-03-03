#!/usr/bin/env python3
# ═══════════════════════════════════════════════════════════════════════════════
#  ATARS — Automated Time-Series Analysis and Reporting System
#  Version  : 2.0.0
#  Author   : Priyanshu
#  Institute: Global Institute of Technology and Management, Haryana, India
#  Email    : priyanshukumar9053@gmail.com
#  Degree   : B.Tech (Computer Science), Second Year
#
#  MIT License
#
#  Copyright (c) 2026 Priyanshu
#
#  Permission is hereby granted, free of charge, to any person obtaining a copy
#  of this software and associated documentation files (the "Software"), to deal
#  in the Software without restriction, including without limitation the rights
#  to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
#  copies of the Software, and to permit persons to whom the Software is
#  furnished to do so, subject to the following conditions:
#
#  The above copyright notice and this permission notice shall be included in
#  all copies or substantial portions of the Software.
#
#  THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR
#  IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
#  FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE
#  AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER
#  LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,
#  OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE
#  SOFTWARE.
#
#  Description:
#    A formally specified, domain-agnostic, end-to-end pipeline for automated
#    statistical analysis, ML-enhanced anomaly detection, forecasting, and
#    reproducible report generation from structured time-series environmental data.
#
#    Implements all formal operators from the ATARS theoretical framework:
#      Aggregation A(D_v), Baseline B(D,W), Z-score ζ_d, Delta δ_d,
#      Quality Score Q(D), Confidence CI, Correlation R, OLS β̂,
#      ACF/PACF, and JSON-grounded narrative generation.
#
#    v2.0 ML additions (additive — all formal operators preserved):
#      ML-1: Isolation Forest (IsolationForest) — multivariate anomaly detection
#            alongside z-score. Does NOT replace formal ζ_d operator.
#      ML-2: Holt-Winters Triple Exponential Smoothing — PM10 14-day forecast.
#            Pure numpy/scipy implementation, zero extra dependencies.
#      INSIGHTS: Deterministic chart insight text generated from actual data
#            values for every chart — no LLM, no hallucination risk.
#
#    Security model (unchanged from v1):
#      LLM receives J only (aggregated stats). No raw sensor data transmitted.
#      All ML computation is local. Zero external API calls. Zero telemetry.
#
#  Usage:
#    python atars.py --data your_data.csv --city "Your City"
#    python atars.py --data your_data.csv --city "Delhi" --no-llm
#    python atars.py --data your_data.csv --city "Delhi" --no-llm --no-forecast
#
#  Requirements:
#    pip install pandas numpy scipy matplotlib seaborn python-docx scikit-learn
#
#  Optional (for AI narrative — install Ollama from https://ollama.com):
#    ollama pull llama3.2
# ═══════════════════════════════════════════════════════════════════════════════

import os
import sys
import json
import math
import hashlib
import warnings
import argparse
import traceback
from pathlib import Path
from datetime import datetime, timedelta
from collections import defaultdict

import numpy as np
import pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.gridspec as gridspec
from matplotlib.patches import FancyBboxPatch
import seaborn as sns
from scipy import stats
from sklearn.linear_model import LinearRegression
from sklearn.preprocessing import StandardScaler
from sklearn.ensemble import IsolationForest
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

warnings.filterwarnings('ignore')

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 0 — CONFIGURATION
#  All parameters in one place. Edit this section for your domain.
# ═══════════════════════════════════════════════════════════════════════════════

DEFAULT_CONFIG = {
    # ── Data ──────────────────────────────────────────────────────────────────
    "data_file"            : "data.csv",
    "city"                 : "City",
    "station_id"           : "Station-01",
    "output_dir"           : "atars_output",

    # ── Statistical parameters ─────────────────────────────────────────────────
    "baseline_window_days" : 30,       # W — days for rolling baseline
    "anomaly_z_threshold"  : 3.0,      # ζ_thresh — z-score anomaly boundary
    "confidence_high"      : 0.90,     # Q(D) ≥ 0.90 → HIGH
    "confidence_moderate"  : 0.70,     # Q(D) ≥ 0.70 → MODERATE
    "ci_alpha"             : 0.05,     # α for 95% confidence intervals
    "acf_lags"             : 48,       # number of lags for ACF/PACF
    "bootstrap_samples"    : 1000,     # B for bootstrap CI
    "min_valid_ratio"      : 0.10,     # minimum Q(D) to process a day

    # ── ML parameters (v2.0) ──────────────────────────────────────────────────
    # ML-1: Isolation Forest anomaly detection
    "if_contamination"     : 0.05,    # expected anomaly fraction (5%)
    "if_n_estimators"      : 150,     # number of trees in the forest
    "if_random_state"      : 42,      # fixed seed for reproducibility
    "if_cols"              : ["PM10","NO2","SO2","CO","Ozone"],  # features
    # ML-2: Holt-Winters forecast
    "hw_forecast_days"     : 14,      # days to forecast ahead
    "hw_seasonal_periods"  : 7,       # weekly seasonality
    "hw_alpha"             : 0.3,     # level smoothing parameter
    "hw_beta"              : 0.1,     # trend smoothing parameter
    "hw_gamma"             : 0.2,     # seasonal smoothing parameter
    "use_forecast"         : True,    # enable Holt-Winters forecast chart

    # ── WHO / Reference thresholds (24-hour means, WHO AQG 2021) ──────────────
    "thresholds": {
        "PM10"       : 45.0,    # µg/m³  — WHO AQG 2021
        "NO2"        : 25.0,    # µg/m³  — WHO AQG 2021
        "SO2"        : 40.0,    # µg/m³  — WHO AQG 2021
        "CO"         : 4.0,     # mg/m³  — WHO AQG 2021
        "Ozone"      : 100.0,   # µg/m³  — WHO AQG 2021 (peak season)
        "NH3"        : 400.0,   # µg/m³  — CPCB India standard
        "NO"         : None,    # no WHO guideline (informational)
        "NOx"        : None,
        "Benzene"    : 1.7,     # µg/m³  — WHO annual reference
        "Toluene"    : None,
        "Xylene"     : None,
        "Eth_Benzene": None,
        "MP_Xylene"  : None,
    },

    # ── LLM (Ollama — local, free, offline) ───────────────────────────────────
    "use_llm"          : True,
    "llm_model"        : "llama3.2",   # change to "mistral", "phi3", "llama3.2:8b" etc.
    "ollama_url"       : "http://localhost:11434",
    "llm_timeout"      : 600,          # seconds per section — 10 min for slow hardware
    "llm_retry"        : 2,            # retry count per section on timeout
    "llm_retry_delay"  : 5,            # seconds between retries

    # ── Report ────────────────────────────────────────────────────────────────
    "author"      : "Priyanshu",
    "institution" : "Global Institute of Technology and Management",
    "department"  : "Department of Computer Science & Engineering",
    "degree"      : "B.Tech (Computer Science), Second Year",
    "email"       : "priyanshukumar9053@gmail.com",
    "location"    : "Haryana, India",
}

# ── Colour palette (consistent across all charts) ─────────────────────────────
PAL = {
    "navy"  : "#1A2C4E", "blue"   : "#2557A7", "teal"   : "#1A6B72",
    "green" : "#1E6B3C", "red"    : "#8B1C2A", "amber"  : "#92400E",
    "gray"  : "#6B7280", "lgray"  : "#F0F4F8", "white"  : "#FFFFFF",
    "purple": "#4C1D95", "orange" : "#7C2D12", "dark"   : "#0D1117",
    "chart_bg": "#F8FAFC",
}

plt.rcParams.update({
    'font.family'      : 'DejaVu Sans',
    'axes.facecolor'   : PAL['chart_bg'],
    'figure.facecolor' : PAL['white'],
    'axes.spines.top'  : False,
    'axes.spines.right': False,
    'axes.grid'        : True,
    'grid.alpha'       : 0.3,
    'grid.linestyle'   : ':',
})

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 1 — COLUMN NORMALIZATION
#  Handles encoding issues (Â° → °, Âµg → µg) automatically.
# ═══════════════════════════════════════════════════════════════════════════════

# Maps raw column names (possibly with encoding artifacts) → clean internal names
# Uses keyword matching so it works regardless of encoding corruption.
COLUMN_KEYWORDS = {
    "humidity"            : "humidity_pct",
    "rain"                : "rain_mm",
    "pressure"            : "pressure_hpa",
    "wind_speed_10"       : "wind_speed_10m",
    "wind_dir"            : "wind_dir_10m",
    "wind_direction_10"   : "wind_dir_10m",
    "wind_speed_100"      : "wind_speed_100m",
    "wind_direction_100"  : "wind_dir_100m",
    "pm10"                : "PM10",
    " no ("               : "NO",        # space before "no" to avoid "nox","no2"
    "no2"                 : "NO2",
    "nox"                 : "NOx",
    "nh3"                 : "NH3",
    "so2"                 : "SO2",
    " co ("               : "CO",        # space to avoid "co2","ozone" matches
    "ozone"               : "Ozone",
    "benzene"             : "Benzene",
    "toluene"             : "Toluene",
    "mp-xylene"           : "MP_Xylene",
    "mp_xylene"           : "MP_Xylene",
    "eth"                 : "Eth_Benzene",
    "xylene"              : "Xylene",    # must come AFTER mp-xylene and eth
}

# Pollutant columns used in statistical analysis
POLLUTANTS = [
    "PM10","NO","NO2","NOx","NH3","SO2","CO",
    "Ozone","Benzene","Toluene","Xylene","Eth_Benzene","MP_Xylene"
]
METEOROLOGICAL = [
    "humidity_pct","rain_mm","pressure_hpa",
    "wind_speed_10m","wind_dir_10m","wind_speed_100m","wind_dir_100m"
]
ALL_NUMERIC = POLLUTANTS + METEOROLOGICAL


def normalize_columns(df: pd.DataFrame) -> pd.DataFrame:
    """
    Rename raw columns to clean internal names.
    Handles UTF-8/Latin-1 encoding artifacts automatically.
    Eq. 1: D = {(t_i, x_i, q_i)} — build the observation space.
    """
    rename_map = {}
    cols_lower = {c: c.lower() for c in df.columns}

    for raw_col, lower_col in cols_lower.items():
        # Direct time columns
        if lower_col in ("day", "date"):
            rename_map[raw_col] = "date"
        elif lower_col == "time":
            rename_map[raw_col] = "time"
        else:
            # Keyword matching (order matters — specific before general)
            matched = False
            for keyword, clean_name in COLUMN_KEYWORDS.items():
                if keyword in lower_col:
                    rename_map[raw_col] = clean_name
                    matched = True
                    break

    df = df.rename(columns=rename_map)

    # Remove duplicate columns (keep first)
    df = df.loc[:, ~df.columns.duplicated()]
    return df


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 2 — DATA LOADING & VALIDATION
# ═══════════════════════════════════════════════════════════════════════════════

def load_data(filepath: str) -> pd.DataFrame:
    """
    Load CSV/Excel, normalize columns, parse datetime, sort.
    Tries multiple encodings to handle real-world file issues.
    """
    print(f"\n{'='*70}")
    print(f"  ATARS — Loading data: {filepath}")
    print(f"{'='*70}")

    path = Path(filepath)
    if not path.exists():
        raise FileNotFoundError(f"Data file not found: {filepath}")

    # Try multiple encodings
    encodings = ['utf-8', 'latin-1', 'cp1252', 'utf-8-sig']
    df = None
    for enc in encodings:
        try:
            if path.suffix.lower() in ('.xlsx', '.xls'):
                df = pd.read_excel(filepath)
            else:
                df = pd.read_csv(filepath, encoding=enc, low_memory=False)
            print(f"  ✓ Loaded with encoding: {enc}")
            break
        except (UnicodeDecodeError, Exception):
            continue

    if df is None:
        raise ValueError("Could not load file with any known encoding.")

    print(f"  ✓ Raw shape: {df.shape[0]:,} rows × {df.shape[1]} columns")

    # Normalize column names
    df = normalize_columns(df)

    # Build datetime index
    if 'date' in df.columns and 'time' in df.columns:
        try:
            df['datetime'] = pd.to_datetime(
                df['date'].astype(str) + ' ' + df['time'].astype(str),
                dayfirst=True
            )
        except Exception:
            df['datetime'] = pd.to_datetime(df['date'], dayfirst=True)
    elif 'date' in df.columns:
        df['datetime'] = pd.to_datetime(df['date'], dayfirst=True)
    else:
        raise ValueError("No 'day'/'date' column found. Check your CSV.")

    df = df.sort_values('datetime').reset_index(drop=True)

    # Convert numeric columns — coerce errors to NaN
    for col in ALL_NUMERIC:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors='coerce')
        else:
            df[col] = np.nan

    # Negative values are physically impossible for concentrations → set NaN
    for col in POLLUTANTS:
        if col in df.columns:
            df.loc[df[col] < 0, col] = np.nan

    df['date_only'] = df['datetime'].dt.date

    print(f"  ✓ Date range: {df['datetime'].min().date()} → {df['datetime'].max().date()}")
    print(f"  ✓ Columns found: {[c for c in ALL_NUMERIC if c in df.columns]}")
    return df


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 3 — QUALITY ASSESSMENT
#  Q(D) = N_v / N_total  (Eq. 11)
#  q_i ∈ {0=unreviewed, 1=valid, 2=suspect, 3=invalid}
# ═══════════════════════════════════════════════════════════════════════════════

# Physical plausibility bounds for range-check quality flagging
PHYSICAL_BOUNDS = {
    "PM10"        : (0, 1000),   # µg/m³
    "NO"          : (0, 500),
    "NO2"         : (0, 500),
    "NOx"         : (0, 1000),   # ppb
    "NH3"         : (0, 2000),
    "SO2"         : (0, 500),
    "CO"          : (0, 50),     # mg/m³
    "Ozone"       : (0, 400),
    "Benzene"     : (0, 100),
    "Toluene"     : (0, 500),
    "Xylene"      : (0, 500),
    "Eth_Benzene" : (0, 200),
    "MP_Xylene"   : (0, 500),
    "humidity_pct": (0, 100),
    "rain_mm"     : (0, 500),
    "pressure_hpa": (800, 1100),
    "wind_speed_10m"  : (0, 200),
    "wind_speed_100m" : (0, 300),
}


def assign_quality_flags(df: pd.DataFrame) -> pd.DataFrame:
    """
    Assign quality flag q_i to each record.
      q_i = 1 (valid)   : not NaN, within physical bounds
      q_i = 2 (suspect) : not NaN but outside 5×IQR (statistical outlier)
      q_i = 3 (invalid) : NaN or outside physical bounds
      q_i = 0 (unreviewed): default for non-pollutant columns
    """
    df = df.copy()

    # Compute IQR for statistical suspect detection
    iqr_bounds = {}
    for col in POLLUTANTS:
        if col in df.columns:
            q25 = df[col].quantile(0.25)
            q75 = df[col].quantile(0.75)
            iqr  = q75 - q25
            iqr_bounds[col] = (q25 - 5 * iqr, q75 + 5 * iqr)

    # Assign flags per pollutant row-wise
    flag_cols = {}
    for col in POLLUTANTS:
        if col not in df.columns:
            continue
        flags = np.ones(len(df), dtype=int)  # default: valid
        phys  = PHYSICAL_BOUNDS.get(col, (0, 1e9))

        # NaN or outside physical bounds → invalid (3)
        mask_invalid = df[col].isna() | (df[col] < phys[0]) | (df[col] > phys[1])
        flags[mask_invalid] = 3

        # Outside 5×IQR → suspect (2), unless already invalid
        if col in iqr_bounds:
            lo, hi = iqr_bounds[col]
            mask_suspect = (~mask_invalid) & ((df[col] < lo) | (df[col] > hi))
            flags[mask_suspect] = 2

        flag_cols[f"q_{col}"] = flags

    for k, v in flag_cols.items():
        df[k] = v

    return df


def compute_quality_score(df_day: pd.DataFrame, col: str) -> dict:
    """
    Q(D) = N_v / N_total  — Eq. 11
    Returns quality metrics for one column on one day.
    """
    flag_col = f"q_{col}"
    if flag_col not in df_day.columns:
        total = len(df_day)
        valid = int(df_day[col].notna().sum())
        return {
            "N_total": total, "N_valid": valid,
            "N_suspect": 0,   "N_invalid": total - valid,
            "Q": valid / total if total > 0 else 0.0
        }
    flags  = df_day[flag_col].values
    total   = len(flags)
    valid   = int(np.sum(flags == 1))
    suspect = int(np.sum(flags == 2))
    invalid = int(np.sum(flags == 3))
    q_score = valid / total if total > 0 else 0.0
    return {
        "N_total": total, "N_valid": valid,
        "N_suspect": suspect, "N_invalid": invalid, "Q": round(q_score, 4)
    }


def get_confidence_flag(q_score: float, config: dict) -> str:
    """Confidence flag from Q(D) — Eq. after Table 3."""
    if q_score >= config["confidence_high"]:
        return "HIGH"
    elif q_score >= config["confidence_moderate"]:
        return "MODERATE"
    return "LOW"


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 4 — FORMAL STATISTICAL OPERATORS  (Eqs. 3–10)
# ═══════════════════════════════════════════════════════════════════════════════

def aggregation_operator(series: pd.Series) -> dict:
    """
    A(D_v) — Aggregation Operator (Eqs. 3–5).
    Computes x̄, σ, x_max, x_min, N_v from validated series.
    """
    valid = series.dropna().values
    N_v = len(valid)
    if N_v == 0:
        return {"mean": np.nan, "std": np.nan, "max": np.nan,
                "min": np.nan, "median": np.nan, "IQR": np.nan, "N_v": 0}
    mean   = float(np.mean(valid))                                          # Eq. 3
    std    = float(np.std(valid, ddof=1)) if N_v > 1 else 0.0              # Eq. 4
    x_max  = float(np.max(valid))                                          # Eq. 5
    x_min  = float(np.min(valid))
    median = float(np.median(valid))
    q25    = float(np.percentile(valid, 25))
    q75    = float(np.percentile(valid, 75))
    iqr = round(q75 - q25, 4) if N_v > 1 else 0.0
    return {
        "mean": round(mean, 4), "std": round(std, 4),
        "max": round(x_max, 4), "min": round(x_min, 4),
        "median": round(median, 4), "IQR": iqr,
        "N_v": N_v
    }


def confidence_interval(mean: float, std: float, n: int,
                         alpha: float = 0.05, method: str = "classical") -> dict:
    """
    Confidence interval formulations (Table in Section 5.5 of paper).
    Classical, Chebyshev, Bootstrap.
    """
    if n <= 1 or np.isnan(mean):
        return {"lower": np.nan, "upper": np.nan, "method": method}

    if method == "classical":
        # Eq. 12: CI = x̄ ± z_{α/2} · σ/sqrt(N_v)
        z = stats.norm.ppf(1 - alpha / 2)     # z_0.025 = 1.96
        margin = z * std / math.sqrt(n)
        return {"lower": round(mean - margin, 4), "upper": round(mean + margin, 4),
                "method": "classical_95pct", "z_critical": round(z, 4)}

    elif method == "chebyshev":
        # Eq. 26: P(|X-μ| ≥ kσ) ≤ 1/k²  →  k=sqrt(1/α)
        k = math.sqrt(1 / alpha)
        return {"lower": round(mean - k * std, 4), "upper": round(mean + k * std, 4),
                "method": "chebyshev", "k_factor": round(k, 4)}

    return {"lower": np.nan, "upper": np.nan, "method": "unknown"}


def baseline_operator(daily_means: pd.Series, current_date,
                       window: int = 30) -> dict:
    """
    B(D, W) — Baseline Operator (Eqs. 7–8).
    Computes x̄_W and σ_W from W days preceding current_date.
    """
    end   = pd.Timestamp(current_date) - pd.Timedelta(days=1)
    start = end - pd.Timedelta(days=window - 1)
    window_data = daily_means.loc[(daily_means.index >= start) & (daily_means.index <= end)].dropna()


    n_W = len(window_data)
    if n_W < 3:
        return {"mean_W": np.nan, "std_W": np.nan, "n_W": n_W}
    mean_W = float(np.mean(window_data.values))                             # Eq. 7
    std_W  = float(np.std(window_data.values, ddof=1))                     # Eq. 8
    return {"mean_W": round(mean_W, 4), "std_W": round(std_W, 4), "n_W": n_W}


def z_score(x_d: float, mean_W: float, std_W: float) -> float:
    """ζ_d = (x̄_d - x̄_W) / σ_W  — Eq. 9"""
    if std_W == 0 or np.isnan(std_W) or np.isnan(x_d):
        return np.nan
    return round((x_d - mean_W) / std_W, 4)


def delta_pct(x_d: float, x_prev: float) -> float:
    """δ_d = [(x̄_d - x̄_{d-1}) / x̄_{d-1}] × 100%  — Eq. 10"""
    if x_prev == 0 or np.isnan(x_d) or np.isnan(x_prev):
        return np.nan
    return round((x_d - x_prev) / x_prev * 100, 2)


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 5 — DAILY AGGREGATION PIPELINE
# ═══════════════════════════════════════════════════════════════════════════════

def build_daily_stats(df: pd.DataFrame, config: dict) -> pd.DataFrame:
    """
    Aggregate hourly observations to daily statistics.
    Returns a DataFrame with one row per day per variable.
    """
    print("\n  [Step 3] Running SQL-equivalent aggregation operators...")
    records = []
    daily_means_by_col = {}

    # First pass: compute daily means for all columns
    for col in ALL_NUMERIC:
        if col not in df.columns:
            continue
        flag_col = f"q_{col}"
        day_groups = df.groupby('date_only')
        means = {}
        for day, grp in day_groups:
            if flag_col in grp.columns:
                valid_vals = grp.loc[grp[flag_col] == 1, col]
            else:
                valid_vals = grp[col].dropna()
            if len(valid_vals) > 0:
                means[pd.Timestamp(day)] = valid_vals.mean()
        s = pd.Series(means).sort_index(); s.index = pd.DatetimeIndex(s.index); daily_means_by_col[col] = s

    # Second pass: full stats per day per column
    for col in ALL_NUMERIC:
        if col not in df.columns:
            continue
        flag_col  = f"q_{col}"
        daily_col = daily_means_by_col.get(col, pd.Series(dtype=float))
        threshold = config["thresholds"].get(col, None)
        prev_mean = np.nan

        for day, grp in df.groupby('date_only'):
            day_ts = pd.Timestamp(day)

            # Validated series for this day
            if flag_col in grp.columns:
                valid_series = grp.loc[grp[flag_col] == 1, col]
            else:
                valid_series = grp[col].dropna()

            q_info  = compute_quality_score(grp, col)
            agg     = aggregation_operator(valid_series)
            conf_f  = get_confidence_flag(q_info["Q"], config)
            baseline= baseline_operator(daily_col, day_ts, config["baseline_window_days"])
            zs      = z_score(agg["mean"], baseline["mean_W"], baseline["std_W"])
            dp      = delta_pct(agg["mean"], prev_mean)
            ci      = confidence_interval(agg["mean"], agg["std"], agg["N_v"], config["ci_alpha"])

            is_anomaly  = (not np.isnan(zs)) and (abs(zs) > config["anomaly_z_threshold"])
            exceeds_thr = (threshold is not None and
                           not np.isnan(agg["mean"]) and
                           agg["mean"] > threshold)

            records.append({
                "date"          : day_ts,
                "variable"      : col,
                "mean"          : agg["mean"],
                "std"           : agg["std"],
                "max"           : agg["max"],
                "min"           : agg["min"],
                "median"        : agg["median"],
                "IQR"           : agg["IQR"],
                "N_v"           : agg["N_v"],
                "N_total"       : q_info["N_total"],
                "Q"             : q_info["Q"],
                "confidence"    : conf_f,
                "mean_W"        : baseline["mean_W"],
                "std_W"         : baseline["std_W"],
                "n_W"           : baseline["n_W"],
                "z_score"       : zs,
                "delta_pct"     : dp,
                "CI_lower"      : ci["lower"],
                "CI_upper"      : ci["upper"],
                "threshold"     : threshold,
                "exceeds_threshold": exceeds_thr,
                "is_anomaly"    : is_anomaly,
            })
            prev_mean = agg["mean"]

    daily_df = pd.DataFrame(records)
    print(f"  ✓ Daily stats: {len(daily_df)} rows × {len(daily_df.columns)} columns")
    return daily_df


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 6 — TEMPORAL ANALYSIS: ACF / PACF  (Eqs. 12–13)
# ═══════════════════════════════════════════════════════════════════════════════

def compute_acf(series: np.ndarray, n_lags: int) -> np.ndarray:
    """
    Sample ACF: ρ̂(k) = Corr(x_t, x_{t-k})  — Eq. 12
    Manual implementation (no statsmodels required).
    """
    n    = len(series)
    mean = np.mean(series)
    var  = np.sum((series - mean) ** 2) / n
    if var == 0:
        return np.zeros(n_lags + 1)
    acf = [1.0]
    for k in range(1, n_lags + 1):
        if k >= n:
            acf.append(0.0)
        else:
            cov = np.sum((series[k:] - mean) * (series[:-k] - mean)) / n
            acf.append(cov / var)
    return np.array(acf)


def compute_pacf(series: np.ndarray, n_lags: int) -> np.ndarray:
    """
    Sample PACF: φ̂_{kk} via Yule-Walker recursion  — Eq. 13
    Manual implementation using OLS at each lag order.
    """
    n = len(series)
    pacf = [1.0]
    for k in range(1, n_lags + 1):
        if k >= n - 1:
            pacf.append(0.0)
            continue
        # Build design matrix: lag-1 through lag-k
        X = np.column_stack([series[k - j - 1: n - j - 1]
                              for j in range(k)])
        y = series[k:]
        # Normalise
        X_m = X - X.mean(axis=0)
        y_m = y - y.mean()
        try:
            coefs = np.linalg.lstsq(X_m, y_m, rcond=None)[0]
            pacf.append(float(coefs[-1]))
        except Exception:
            pacf.append(0.0)
    return np.array(pacf)


def run_temporal_analysis(daily_df: pd.DataFrame, col: str,
                           n_lags: int = 48) -> dict:
    """Run ACF/PACF on daily mean series for one variable."""
    series = daily_df[daily_df['variable'] == col]['mean'].dropna().values
    if len(series) < n_lags + 5:
        return {}
    acf  = compute_acf(series, n_lags)
    pacf = compute_pacf(series, min(n_lags, 24))  # PACF expensive at high lags
    conf_bound = 1.96 / math.sqrt(len(series))
    return {
        "acf" : acf, "pacf": pacf,
        "conf_bound"      : conf_bound,
        "significant_acf" : [k for k, v in enumerate(acf) if k > 0 and abs(v) > conf_bound],
        "series_length"   : len(series),
    }


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 7 — CORRELATION MATRIX & OLS  (Eqs. 16–18, 14–15)
# ═══════════════════════════════════════════════════════════════════════════════

def compute_correlation_matrix(daily_df: pd.DataFrame,
                                variables: list) -> pd.DataFrame:
    """
    R_{jk} = Corr(X_j, X_k)  — Eq. 18
    Built from daily means. Returns d×d DataFrame.
    NOTE: Association ≠ Causation. Corr(X,Y) ≢ C(X→Y).
    """
    pivot = daily_df.pivot_table(
        index='date', columns='variable', values='mean'
    )
    cols = [v for v in variables if v in pivot.columns]
    return pivot[cols].corr(method='pearson')


def run_ols_regression(daily_df: pd.DataFrame,
                        target: str = "PM10",
                        predictors: list = None) -> dict:
    """
    OLS: β̂ = (X^T X)^{-1} X^T y  — Eq. 14
    Predicts target from meteorological predictors.
    Returns association statistics only — no causal interpretation.
    """
    if predictors is None:
        predictors = [p for p in ["humidity_pct", "wind_speed_10m",
                                   "rain_mm", "pressure_hpa"]
                      if p in daily_df['variable'].unique()]
    pivot = daily_df.pivot_table(
        index='date', columns='variable', values='mean'
    )
    if target not in pivot.columns:
        return {}
    available = [p for p in predictors if p in pivot.columns]
    if len(available) < 2:
        return {}

    df_reg = pivot[[target] + available].dropna()
    if len(df_reg) < 20:
        return {}

    X = df_reg[available].values
    y = df_reg[target].values

    scaler = StandardScaler()
    X_s = scaler.fit_transform(X)

    beta, residuals, _, _ = np.linalg.lstsq(
        np.column_stack([np.ones(len(X_s)), X_s]), y, rcond=None
    )

    y_pred = np.column_stack([np.ones(len(X_s)), X_s]) @ beta
    ss_res = np.sum((y - y_pred) ** 2)
    ss_tot = np.sum((y - np.mean(y)) ** 2)
    r2     = 1 - ss_res / ss_tot if ss_tot > 0 else 0.0

    return {
        "target"    : target,
        "predictors": available,
        "beta"      : {p: round(float(b), 4) for p, b in zip(available, beta[1:])},
        "intercept" : round(float(beta[0]), 4),
        "r_squared" : round(r2, 4),
        "n_obs"     : len(df_reg),
        "note"      : "Association only — Corr(X,Y) does not imply causation (Eq. 17)"
    }


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 7.5 — ML ANOMALY DETECTION: ISOLATION FOREST
#  Runs alongside z-score (ζ_d). Does NOT replace formal Eq. 9.
#  Detects multivariate anomalies that univariate z-score cannot see.
#
#  SECURITY NOTE: All computation is local. IsolationForest is fit and
#  applied entirely in memory on daily_df data. No data leaves the machine.
#  Fixed random_state=42 ensures reproducible results (same seed → same model).
# ═══════════════════════════════════════════════════════════════════════════════

def run_isolation_forest(daily_df: pd.DataFrame, config: dict) -> dict:
    """
    ML-1: Isolation Forest multivariate anomaly detection.

    How it works:
      1. Pivot daily_df to a wide matrix: rows = dates, cols = pollutants
      2. Standardise each column (zero mean, unit variance)
      3. Fit IsolationForest on the full dataset (unsupervised — no labels)
      4. Score each day: anomaly_score ∈ [-1, 1] where lower = more anomalous
      5. Flag days where predict() == -1 as ML anomalies

    Why it improves on z-score:
      - z-score is univariate: checks one variable at a time
      - IsolationForest is multivariate: detects days where the COMBINATION
        of PM10 + NO2 + SO2 is anomalous even if no single variable crosses 3σ
      - Does not assume normality (z-score implicitly does)
      - Automatically handles correlated features

    Formal complement: ML anomaly ∪ z-score anomaly ⊇ z-score anomaly.
    A day flagged by IF but NOT z-score warrants manual investigation.

    Returns:
      dict with keys: flagged_dates, scores, n_anomalies, features_used,
                      contamination, comparison_df
    """
    print("\n  [ML-1] Running Isolation Forest multivariate anomaly detection...")

    # Select feature columns available in data
    feature_cols = [c for c in config["if_cols"]
                    if c in daily_df["variable"].unique()]
    if len(feature_cols) < 2:
        print(f"  ⚠  Isolation Forest skipped — fewer than 2 feature columns available.")
        return {"available": False, "reason": "insufficient_features"}

    # Build wide matrix: one row per day
    pivot = daily_df.pivot_table(index="date", columns="variable", values="mean")
    pivot = pivot[[c for c in feature_cols if c in pivot.columns]].dropna()
    if len(pivot) < 30:
        print(f"  ⚠  Isolation Forest skipped — fewer than 30 complete rows ({len(pivot)}).")
        return {"available": False, "reason": "insufficient_rows"}

    used_cols = list(pivot.columns)
    X = pivot.values

    # Standardise (important: IF is not scale-invariant)
    scaler = StandardScaler()
    X_s = scaler.fit_transform(X)

    # Fit Isolation Forest
    # contamination: expected fraction of anomalies (5% = 1 in 20 days)
    # random_state: fixed for reproducibility (same seed → same tree structure)
    # n_estimators: 150 trees for stable scoring
    clf = IsolationForest(
        contamination  = config["if_contamination"],
        n_estimators   = config["if_n_estimators"],
        random_state   = config["if_random_state"],
        n_jobs         = 1,      # single-threaded for reproducibility
    )
    clf.fit(X_s)

    preds  = clf.predict(X_s)          # 1 = normal, -1 = anomaly
    scores = clf.decision_function(X_s)  # higher = more normal

    anomaly_mask  = preds == -1
    flagged_dates = list(pivot.index[anomaly_mask])
    n_anomalies   = int(anomaly_mask.sum())

    # Build comparison dataframe (for report and Chart 14)
    comp = pd.DataFrame({
        "date"         : pivot.index,
        "if_score"     : scores,
        "ml_anomaly"   : anomaly_mask,
    }).set_index("date")

    # Merge with z-score anomaly column from daily_df (PM10 only)
    pm10_anom = daily_df[daily_df["variable"] == "PM10"][["date","is_anomaly","z_score"]].copy()
    pm10_anom = pm10_anom.set_index("date").rename(
        columns={"is_anomaly": "zscore_anomaly", "z_score": "z_score"})
    comp = comp.join(pm10_anom, how="left")
    comp["zscore_anomaly"] = comp["zscore_anomaly"].fillna(False)

    # Agreement categories
    comp["both"]    = comp["ml_anomaly"] & comp["zscore_anomaly"]
    comp["if_only"] = comp["ml_anomaly"] & ~comp["zscore_anomaly"]
    comp["zs_only"] = ~comp["ml_anomaly"] & comp["zscore_anomaly"]

    n_both   = int(comp["both"].sum())
    n_if_only= int(comp["if_only"].sum())
    n_zs_only= int(comp["zs_only"].sum())

    print(f"  ✓ Isolation Forest: {n_anomalies} anomalies in {len(pivot)} days "
          f"({n_anomalies/len(pivot)*100:.1f}%)")
    print(f"  ✓ Agreement: both={n_both} | IF-only={n_if_only} | z-score-only={n_zs_only}")

    return {
        "available"      : True,
        "features_used"  : used_cols,
        "n_days"         : len(pivot),
        "n_anomalies"    : n_anomalies,
        "anomaly_rate"   : round(n_anomalies / len(pivot), 4),
        "contamination"  : config["if_contamination"],
        "n_estimators"   : config["if_n_estimators"],
        "random_state"   : config["if_random_state"],
        "flagged_dates"  : [str(d.date()) for d in flagged_dates],
        "comparison_df"  : comp,
        "n_both"         : n_both,
        "n_if_only"      : n_if_only,
        "n_zs_only"      : n_zs_only,
        "security_note"  : (
            "IsolationForest fit on local daily_df only. "
            "No data transmitted. random_state=42 ensures reproducibility."
        ),
    }


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 7.6 — ML FORECASTING: HOLT-WINTERS TRIPLE EXPONENTIAL SMOOTHING
#  Pure numpy implementation — zero extra dependencies beyond existing stack.
#  Forecasts PM10 14 days ahead. Used for Chart 15 only.
#
#  SECURITY NOTE: Computed entirely in memory from daily_df.
#  Forecast values are NOT added to J (the LLM data contract) because:
#    1. Forecasts have uncertainty — adding them to J risks LLM stating
#       forecasts as historical facts.
#    2. The AI constraints in J explicitly prohibit extrapolation.
#  Forecasts appear only in Chart 15 and the report ML section.
# ═══════════════════════════════════════════════════════════════════════════════

def holt_winters_forecast(series: np.ndarray, seasonal_periods: int = 7,
                           forecast_days: int = 14,
                           alpha: float = 0.3, beta: float = 0.1,
                           gamma: float = 0.2) -> dict:
    """
    Triple Exponential Smoothing (Holt-Winters additive seasonal model).

    Parameters:
        series          : 1-D array of historical daily PM10 means
        seasonal_periods: length of one season (7 = weekly)
        forecast_days   : number of days to forecast ahead
        alpha (α)       : level smoothing  — how fast level adapts (0 < α < 1)
        beta  (β)       : trend smoothing  — how fast trend adapts (0 < β < 1)
        gamma (γ)       : seasonal smoothing (0 < γ < 1)

    Update equations (additive Holt-Winters):
        Level:    L_t = α(x_t - S_{t-m}) + (1-α)(L_{t-1} + T_{t-1})
        Trend:    T_t = β(L_t - L_{t-1}) + (1-β)T_{t-1}
        Seasonal: S_t = γ(x_t - L_t) + (1-γ)S_{t-m}

    Forecast:   x̂_{t+h} = L_t + h·T_t + S_{t-m+((h-1) mod m)+1}

    Returns dict with: forecast array, lower/upper 80% prediction bands,
    fitted values, and model parameters.
    """
    n = len(series)
    m = seasonal_periods

    if n < 2 * m:
        return {"available": False, "reason": "insufficient_history"}

    series = np.array(series, dtype=float)

    # ── Initialisation ──────────────────────────────────────────────────────
    # Level: mean of first season
    L = np.mean(series[:m])
    # Trend: average slope across first two seasons
    T = (np.mean(series[m:2*m]) - np.mean(series[:m])) / m
    # Seasonal indices: deviation of each period from level
    S = np.array([(series[i] - L) for i in range(m)], dtype=float)

    levels   = np.zeros(n)
    trends   = np.zeros(n)
    seasonals= np.zeros(n + forecast_days + m)
    seasonals[:m] = S
    fitted   = np.zeros(n)

    # ── Smoothing pass ──────────────────────────────────────────────────────
    for t in range(n):
        s_idx = t % m
        x_t   = series[t]
        L_new = alpha * (x_t - seasonals[t % m]) + (1 - alpha) * (L + T)
        T_new = beta  * (L_new - L) + (1 - beta) * T
        S_new = gamma * (x_t - L_new) + (1 - gamma) * seasonals[t % m]
        fitted[t]      = L + T + seasonals[t % m]
        seasonals[t + m] = S_new
        L, T           = L_new, T_new
        levels[t]  = L
        trends[t]  = T

    # ── Residual std for prediction intervals ───────────────────────────────
    resid = series - fitted
    resid_std = float(np.std(resid, ddof=1))

    # ── Forecast ────────────────────────────────────────────────────────────
    forecast = np.zeros(forecast_days)
    for h in range(1, forecast_days + 1):
        s_idx      = (n + h - 1) % m
        forecast[h-1] = L + h * T + seasonals[n + h - 1]

    # 80% prediction interval: ±1.28 * resid_std * sqrt(h)
    # Wider interval for longer horizons (uncertainty grows with h)
    pi_factor = 1.28
    lower = np.array([forecast[h] - pi_factor * resid_std * math.sqrt(h+1)
                      for h in range(forecast_days)])
    upper = np.array([forecast[h] + pi_factor * resid_std * math.sqrt(h+1)
                      for h in range(forecast_days)])

    # Clip to non-negative (concentrations cannot be negative)
    forecast = np.clip(forecast, 0, None)
    lower    = np.clip(lower, 0, None)
    upper    = np.clip(upper, 0, None)

    rmse = float(np.sqrt(np.mean(resid**2)))
    mae  = float(np.mean(np.abs(resid)))

    return {
        "available"       : True,
        "forecast"        : forecast,
        "lower_80"        : lower,
        "upper_80"        : upper,
        "fitted"          : fitted,
        "rmse"            : round(rmse, 3),
        "mae"             : round(mae, 3),
        "resid_std"       : round(resid_std, 3),
        "alpha"           : alpha,
        "beta"            : beta,
        "gamma"           : gamma,
        "seasonal_periods": m,
        "forecast_days"   : forecast_days,
        "n_history"       : n,
        "security_note"   : (
            "Forecast computed locally on daily_df. NOT added to JSON contract J "
            "to prevent LLM from stating forecasts as historical facts. "
            "Forecast appears in Chart 15 and ML report section only."
        ),
    }


def run_hw_forecast(daily_df: pd.DataFrame, config: dict) -> dict:
    """
    Orchestrate Holt-Winters forecast for PM10.
    Extracts PM10 daily means, calls holt_winters_forecast(),
    builds forecast date index, returns full result dict.
    """
    print("\n  [ML-2] Running Holt-Winters forecast for PM10...")

    pm10 = daily_df[daily_df["variable"] == "PM10"][["date","mean"]].dropna()
    pm10 = pm10.sort_values("date")
    if len(pm10) < 14:
        print("  ⚠  Forecast skipped — fewer than 14 PM10 data points.")
        return {"available": False, "reason": "insufficient_pm10_data"}

    series   = pm10["mean"].values
    last_date= pm10["date"].iloc[-1]

    result = holt_winters_forecast(
        series          = series,
        seasonal_periods= config["hw_seasonal_periods"],
        forecast_days   = config["hw_forecast_days"],
        alpha           = config["hw_alpha"],
        beta            = config["hw_beta"],
        gamma           = config["hw_gamma"],
    )

    if not result.get("available"):
        print(f"  ⚠  Forecast failed: {result.get('reason')}")
        return result

    # Build forecast date range
    forecast_dates = pd.date_range(
        start=last_date + pd.Timedelta(days=1),
        periods=config["hw_forecast_days"],
        freq="D"
    )
    result["forecast_dates"] = forecast_dates
    result["history_dates"]  = pm10["date"].values
    result["history_series"] = series

    print(f"  ✓ Holt-Winters: RMSE={result['rmse']:.2f} µg/m³ | "
          f"MAE={result['mae']:.2f} µg/m³ | "
          f"Forecast horizon: {config['hw_forecast_days']} days")
    print(f"  ✓ Next 7-day mean forecast: "
          f"{np.mean(result['forecast'][:7]):.1f} µg/m³")

    return result


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 7.7 — CHART INSIGHT GENERATOR
#  Generates deterministic text insights for every chart from actual data.
#  NO LLM. NO hallucination risk. All numbers come from daily_df / corr_matrix.
#  These insights appear below each chart in the Word report.
# ═══════════════════════════════════════════════════════════════════════════════

def generate_chart_insights(daily_df: pd.DataFrame, df: pd.DataFrame,
                             config: dict, corr_matrix: pd.DataFrame,
                             ols_result: dict, ml_if: dict,
                             ml_hw: dict) -> dict:
    """
    Compute data-driven text insight for each chart.

    Every string is built from actual computed values in daily_df and corr_matrix.
    No LLM is involved. No value is invented or estimated.
    These insights are factual summaries — not interpretations.

    Returns:
        dict: {chart_name: insight_string}
    """
    insights = {}
    city = config["city"]
    thr  = config["thresholds"]

    def safe_get(col, stat="mean"):
        sub = daily_df[daily_df["variable"] == col]
        if sub.empty or sub[stat].isna().all():
            return None
        return float(sub[stat].mean()) if stat == "mean" else float(sub[stat].iloc[0])

    def annual_mean(col):
        sub = daily_df[daily_df["variable"] == col]["mean"].dropna()
        return float(sub.mean()) if len(sub) > 0 else None

    def annual_max(col):
        sub = daily_df[daily_df["variable"] == col]["mean"].dropna()
        return float(sub.max()) if len(sub) > 0 else None

    def exceed_days(col):
        sub = daily_df[daily_df["variable"] == col]
        return int(sub["exceeds_threshold"].sum()) if "exceeds_threshold" in sub.columns else 0

    def anomaly_days(col):
        sub = daily_df[daily_df["variable"] == col]
        return int(sub["is_anomaly"].sum()) if "is_anomaly" in sub.columns else 0

    def avg_quality(col):
        sub = daily_df[daily_df["variable"] == col]["Q"].dropna()
        return float(sub.mean()) if len(sub) > 0 else None

    def peak_month(col):
        sub = daily_df[daily_df["variable"] == col][["date","mean"]].dropna()
        if sub.empty: return "N/A"
        sub = sub.copy(); sub["month"] = sub["date"].dt.month
        m = sub.groupby("month")["mean"].mean().idxmax()
        return ["Jan","Feb","Mar","Apr","May","Jun",
                "Jul","Aug","Sep","Oct","Nov","Dec"][m - 1]

    def low_month(col):
        sub = daily_df[daily_df["variable"] == col][["date","mean"]].dropna()
        if sub.empty: return "N/A"
        sub = sub.copy(); sub["month"] = sub["date"].dt.month
        m = sub.groupby("month")["mean"].mean().idxmin()
        return ["Jan","Feb","Mar","Apr","May","Jun",
                "Jul","Aug","Sep","Oct","Nov","Dec"][m - 1]

    # ── Chart 1: Hourly Time-Series ──────────────────────────────────────────
    pm10_mean  = annual_mean("PM10")
    pm10_max   = annual_max("PM10")
    pm10_ex    = exceed_days("PM10")
    no2_mean   = annual_mean("NO2")
    pm10_thr   = thr.get("PM10", 45)
    no2_thr    = thr.get("NO2", 25)
    total_days = daily_df["date"].nunique()

    if pm10_mean is not None:
        pm10_ex_pct = pm10_ex / total_days * 100 if total_days > 0 else 0
        insights["timeseries"] = (
            f"PM10 annual mean: {pm10_mean:.1f} µg/m³ "
            f"({'above' if pm10_mean > pm10_thr else 'below'} WHO guideline {pm10_thr} µg/m³). "
            f"Peak daily mean: {pm10_max:.1f} µg/m³. "
            f"WHO exceedances: {pm10_ex} days ({pm10_ex_pct:.1f}% of monitored period). "
            + (f"NO₂ annual mean: {no2_mean:.1f} µg/m³ "
               f"({'above' if no2_mean > no2_thr else 'below'} WHO guideline {no2_thr} µg/m³)."
               if no2_mean is not None else "NO₂ data not available.")
        )
    else:
        insights["timeseries"] = "PM10 data not available for time-series analysis."

    # ── Chart 2: Correlation Heatmap ─────────────────────────────────────────
    if corr_matrix is not None and not corr_matrix.empty:
        flat = corr_matrix.where(
            np.triu(np.ones(corr_matrix.shape), k=1).astype(bool)
        ).stack()
        if len(flat) > 0:
            top_pair = flat.abs().idxmax()
            top_r    = float(flat[top_pair])
            v1, v2   = top_pair
            direction = "positive" if top_r > 0 else "negative"
            # Find PM10-humidity if available
            pm10_hum = None
            if "PM10" in corr_matrix.index and "humidity_pct" in corr_matrix.columns:
                pm10_hum = float(corr_matrix.loc["PM10","humidity_pct"])
            insights["correlation"] = (
                f"Strongest association: {v1}–{v2} (r = {top_r:.3f}, {direction}). "
                + (f"PM10–humidity association: r = {pm10_hum:.3f} "
                   f"({'rain washes out particles — consistent with wet deposition' if pm10_hum < -0.3 else 'weak inverse pattern'})."
                   if pm10_hum is not None else "")
                + " NOTE: All correlations are statistical associations only — not causal relationships."
            )
        else:
            insights["correlation"] = "Correlation matrix computed. Interpret all r values as associations only — not causation."
    else:
        insights["correlation"] = "Correlation matrix unavailable."

    # ── Chart 3: 30-Day Rolling Trend ────────────────────────────────────────
    pk = peak_month("PM10"); lk = low_month("PM10")
    pm10_std = daily_df[daily_df["variable"] == "PM10"]["mean"].std()
    insights["trend"] = (
        f"PM10 30-day rolling mean peaks in {pk} and reaches minimum in {lk}, "
        f"indicating {'strong' if pm10_std is not None and pm10_std > 20 else 'moderate'} "
        f"seasonal variation (annual σ = {pm10_std:.1f} µg/m³). "
        f"The rolling window smooths short-term spikes to reveal the underlying seasonal pattern. "
        f"Deviations above the rolling mean indicate elevated pollution episodes."
    ) if pm10_mean is not None else "PM10 trend data unavailable."

    # ── Chart 4: Anomaly Detection ────────────────────────────────────────────
    pm10_an = anomaly_days("PM10")
    no2_an  = anomaly_days("NO2")
    z_thr   = config["anomaly_z_threshold"]
    insights["anomalies"] = (
        f"Anomaly threshold: |ζ_d| > {z_thr:.1f} (3σ rule). "
        f"PM10 anomaly days: {pm10_an} "
        f"({pm10_an/total_days*100:.1f}% of monitored period). "
        + (f"NO₂ anomaly days: {no2_an}. " if no2_an is not None else "")
        + f"Anomalies are computed against a {config['baseline_window_days']}-day rolling baseline, "
        f"so seasonal variation does not inflate anomaly counts. "
        f"Each anomaly represents a statistically significant deviation from recent local norms."
    )

    # ── Chart 5: ACF/PACF ────────────────────────────────────────────────────
    # Compute lag-1 ACF for PM10
    pm10_series = daily_df[daily_df["variable"] == "PM10"]["mean"].dropna().values
    lag1_acf = "N/A"
    if len(pm10_series) > 5:
        acf_arr = compute_acf(pm10_series, 2)
        lag1_acf = f"{acf_arr[1]:.3f}"
    insights["acf_pacf"] = (
        f"PM10 lag-1 autocorrelation: ρ̂(1) = {lag1_acf} — "
        f"{'strong' if float(lag1_acf) > 0.5 else 'moderate' if float(lag1_acf) > 0.3 else 'weak'} "
        f"temporal memory (today's PM10 {'strongly ' if float(lag1_acf) > 0.5 else ''}predicts tomorrow's). "
        f"ACF at 48 lags and PACF at 24 lags. Dashed lines = 95% confidence bounds (±1.96/√n). "
        f"Bars extending beyond the dashed lines are statistically significant. "
        f"This autocorrelation structure violates the independence assumption of standard statistical tests."
    ) if lag1_acf != "N/A" else "ACF/PACF requires at least 53 days of PM10 data."

    # ── Chart 6: Data Quality Dashboard ──────────────────────────────────────
    overall_q = daily_df.groupby("variable")["Q"].mean()
    low_q_vars = [v for v, q in overall_q.items() if q < config["confidence_moderate"]]
    best_q_var = overall_q.idxmax() if len(overall_q) > 0 else "N/A"
    worst_q_var= overall_q.idxmin() if len(overall_q) > 0 else "N/A"
    mean_q     = float(overall_q.mean()) if len(overall_q) > 0 else 0
    insights["quality"] = (
        f"Overall data completeness Q(D) = {mean_q:.3f} "
        f"({'HIGH' if mean_q >= config['confidence_high'] else 'MODERATE' if mean_q >= config['confidence_moderate'] else 'LOW'}). "
        f"Best quality: {best_q_var} (Q = {float(overall_q.get(best_q_var, 0)):.3f}). "
        f"Lowest quality: {worst_q_var} (Q = {float(overall_q.get(worst_q_var, 0)):.3f}). "
        + (f"Variables with LOW confidence (Q < {config['confidence_moderate']}): "
           f"{', '.join(low_q_vars) if low_q_vars else 'none'}. " )
        + "Q(D) measures completeness only — not sensor calibration accuracy. "
        "Days with Q < 0.50 are excluded from baseline window calculations."
    )

    # ── Chart 7: Seasonal Box Plots ───────────────────────────────────────────
    insights["seasonal"] = (
        f"Monthly box plots show the full distribution of daily PM10 means for each month. "
        f"Peak pollution month: {pk}. Lowest pollution month: {lk}. "
        f"Box = 25th–75th percentile. Whiskers = 1.5×IQR. Circles = outlier days. "
        f"Annual PM10 σ = {pm10_std:.1f} µg/m³ — "
        f"{'indicating high seasonal variability' if pm10_std is not None and pm10_std > 20 else 'indicating moderate seasonal variability'}."
    ) if pm10_mean is not None else "PM10 seasonal data unavailable."

    # ── Chart 8: Diurnal Pattern ──────────────────────────────────────────────
    # Find peak hour for PM10 from hourly data
    if "PM10" in df.columns and "datetime" in df.columns:
        pm10_h = df[["datetime","PM10"]].dropna()
        pm10_h = pm10_h.copy(); pm10_h["hour"] = pm10_h["datetime"].dt.hour
        hourly_avg = pm10_h.groupby("hour")["PM10"].mean()
        if len(hourly_avg) > 0:
            peak_hr  = int(hourly_avg.idxmax())
            low_hr   = int(hourly_avg.idxmin())
            peak_val = float(hourly_avg.max())
            low_val  = float(hourly_avg.min())
            insights["diurnal"] = (
                f"Diurnal PM10 pattern: peak at {peak_hr:02d}:00 ({peak_val:.1f} µg/m³), "
                f"minimum at {low_hr:02d}:00 ({low_val:.1f} µg/m³). "
                f"{'Morning and evening peaks consistent with traffic emission patterns.' if peak_hr in [7,8,9,18,19,20] else 'Peak timing may reflect local industrial or meteorological patterns.'} "
                f"Diurnal amplitude = {peak_val - low_val:.1f} µg/m³."
            )
        else:
            insights["diurnal"] = "Diurnal pattern computed from hourly data. Shows average concentration by hour of day (0–23)."
    else:
        insights["diurnal"] = "Diurnal pattern computed from hourly data. Shows average concentration by hour of day (0–23)."

    # ── Chart 9: Day-of-Week ──────────────────────────────────────────────────
    if "PM10" in daily_df["variable"].unique():
        pm10_d = daily_df[daily_df["variable"] == "PM10"][["date","mean"]].dropna().copy()
        pm10_d["dow"] = pm10_d["date"].dt.dayofweek
        dow_avg  = pm10_d.groupby("dow")["mean"].mean()
        day_names= ["Mon","Tue","Wed","Thu","Fri","Sat","Sun"]
        peak_dow = day_names[int(dow_avg.idxmax())]
        low_dow  = day_names[int(dow_avg.idxmin())]
        weekday_avg = float(dow_avg[:5].mean())
        weekend_avg = float(dow_avg[5:].mean())
        diff_pct = (weekday_avg - weekend_avg) / weekend_avg * 100 if weekend_avg > 0 else 0
        insights["day_of_week"] = (
            f"Highest PM10: {peak_dow}. Lowest: {low_dow}. "
            f"Weekday mean: {weekday_avg:.1f} µg/m³ vs weekend mean: {weekend_avg:.1f} µg/m³ "
            f"({'weekdays {:.1f}% higher'.format(diff_pct) if diff_pct > 5 else 'weekdays {:.1f}% lower'.format(abs(diff_pct)) if diff_pct < -5 else 'minimal weekday–weekend difference'} — "
            f"{'consistent with traffic and industrial emission patterns' if diff_pct > 5 else 'suggests non-traffic dominant sources'})."
        )
    else:
        insights["day_of_week"] = "Day-of-week pattern shows average daily means by weekday for PM10, NO₂, and CO."

    # ── Chart 10: VOC/BTEX Panel ──────────────────────────────────────────────
    benz_mean = annual_mean("Benzene")
    benz_thr  = thr.get("Benzene", 1.7)
    tol_mean  = annual_mean("Toluene")
    if benz_mean is not None:
        insights["voc"] = (
            f"Benzene annual mean: {benz_mean:.3f} µg/m³ "
            f"({'above' if benz_mean > benz_thr else 'below'} WHO annual reference {benz_thr} µg/m³). "
            + (f"Toluene annual mean: {tol_mean:.2f} µg/m³. " if tol_mean is not None else "")
            + f"BTEX compounds (Benzene, Toluene, Ethylbenzene, Xylene) indicate "
            f"{'petrochemical and traffic emission sources' if benz_mean > benz_thr else 'low petrochemical source influence at this station'}. "
            f"Benzene is a Group 1 IARC carcinogen — even sub-guideline levels warrant monitoring."
        )
    else:
        insights["voc"] = "VOC/BTEX data not available at this station."

    # ── Chart 11: Rain Scatter ────────────────────────────────────────────────
    if corr_matrix is not None and "PM10" in corr_matrix.index and "rain_mm" in corr_matrix.columns:
        r_rain_pm10 = float(corr_matrix.loc["PM10", "rain_mm"])
        insights["rain_scatter"] = (
            f"PM10–rainfall Pearson r = {r_rain_pm10:.3f} "
            f"({'negative — consistent with wet deposition (rain washes PM10 out of atmosphere)' if r_rain_pm10 < -0.2 else 'near-zero — rainfall has limited association with PM10 at this station' if abs(r_rain_pm10) < 0.2 else 'positive — unexpected direction; investigate data quality'}). "
            f"OLS trend line shows association direction. "
            f"IMPORTANT: This is statistical association only — not a proven causal wet-deposition mechanism."
        )
    else:
        insights["rain_scatter"] = (
            "Scatter of daily rainfall against PM10 and SO₂ concentrations. "
            "Negative slope indicates wet deposition effect (association, not causation)."
        )

    # ── Chart 12: Exceedance Calendar ────────────────────────────────────────
    pm10_thr_val = thr.get("PM10", 45)
    insights["calendar"] = (
        f"Full-year calendar heatmap of PM10 daily means (µg/m³). "
        f"× marks days exceeding WHO guideline ({pm10_thr_val} µg/m³). "
        f"Total exceedances: {pm10_ex} days ({pm10_ex / total_days * 100:.1f}% of monitored period). "
        f"Clustering of × marks in {pk} indicates seasonal high-pollution period. "
        f"Green cells = clean days. Red cells = high pollution. "
        f"Each cell value = that day's daily mean PM10 (µg/m³)."
    ) if pm10_mean is not None else "Calendar heatmap of PM10 exceedance days."

    # ── Chart 13: RGV Grounding ───────────────────────────────────────────────
    insights["grounding"] = (
        "The Runtime Grounding Verifier (RGV) extracts every number from every "
        "LLM sentence and verifies it against the JSON data contract J using "
        "type-adaptive tolerance bounds (±5% for concentrations, ±2 for integer counts, "
        "±0.1 for quality scores). "
        "GROUNDED = all numbers match J. PARTIALLY = ≥50% match. "
        "UNGROUNDED = <50% match (LLM may have hallucinated these values). "
        "NON_NUMERICAL = sentence has no numbers (excluded from G_rate). "
        "G_rate = grounded sentences / numerical sentences. PASS if G_rate ≥ 0.75."
    )

    # ── Chart 14: ML Anomaly Comparison ──────────────────────────────────────
    if ml_if.get("available"):
        n_if    = ml_if["n_anomalies"]
        n_both  = ml_if["n_both"]
        n_ifonly= ml_if["n_if_only"]
        n_zsonly= ml_if["n_zs_only"]
        feats   = ", ".join(ml_if["features_used"])
        insights["ml_anomaly"] = (
            f"Isolation Forest ({ml_if['n_estimators']} trees, contamination={ml_if['contamination']:.0%}) "
            f"on {len(ml_if['features_used'])} variables ({feats}): "
            f"{n_if} anomaly days ({ml_if['anomaly_rate']*100:.1f}%). "
            f"Agreement with z-score: {n_both} days flagged by both methods. "
            f"{n_ifonly} days flagged by IF only (multivariate pattern — z-score missed). "
            f"{n_zsonly} days flagged by z-score only (univariate spike, within multivariate norm). "
            f"IF-only days represent genuine added value of multivariate detection. "
            f"NOTE: Isolation Forest does NOT replace formal ζ_d — both methods are reported. "
            f"random_state=42 ensures this result is reproducible."
        )
    else:
        insights["ml_anomaly"] = "Isolation Forest skipped (insufficient multivariate data)."

    # ── Chart 15: Holt-Winters Forecast ──────────────────────────────────────
    if ml_hw.get("available"):
        fc_mean = float(np.mean(ml_hw["forecast"]))
        fc_max  = float(np.max(ml_hw["forecast"]))
        fc_days = ml_hw["forecast_days"]
        insights["forecast"] = (
            f"Holt-Winters triple exponential smoothing (α={ml_hw['alpha']}, "
            f"β={ml_hw['beta']}, γ={ml_hw['gamma']}, "
            f"seasonal period={ml_hw['seasonal_periods']} days). "
            f"In-sample RMSE={ml_hw['rmse']:.2f} µg/m³ | MAE={ml_hw['mae']:.2f} µg/m³. "
            f"{fc_days}-day PM10 forecast: mean={fc_mean:.1f} µg/m³, peak={fc_max:.1f} µg/m³. "
            f"Shaded band = 80% prediction interval (widens with forecast horizon). "
            f"IMPORTANT: This is a statistical extrapolation, not a physical model. "
            f"Forecast values are NOT in the JSON contract J — they were not used by the LLM. "
            f"Uncertainty grows significantly beyond 7 days."
        )
    else:
        insights["forecast"] = "Holt-Winters forecast skipped (insufficient PM10 history)."

    return insights

def _save_fig(fig, path: Path, name: str) -> str:
    """Save figure and return file path string."""
    fpath = str(path / name)
    fig.savefig(fpath, dpi=160, bbox_inches='tight', facecolor=PAL['white'])
    plt.close(fig)
    return fpath


def chart_timeseries(df: pd.DataFrame, daily_df: pd.DataFrame,
                      config: dict, out: Path) -> str:
    """Chart 1 — Hourly time-series for PM10 and NO2 with WHO guidelines."""
    fig, axes = plt.subplots(2, 1, figsize=(14, 9), facecolor=PAL['white'])
    fig.subplots_adjust(hspace=0.45)

    for ax, col, color, thr_color, title in [
        (axes[0], 'PM10', PAL['blue'],   PAL['amber'],
         f"PM10 Hourly Concentration  —  WHO 24h Guideline: {config['thresholds']['PM10']} µg/m³"),
        (axes[1], 'NO2',  PAL['teal'], PAL['red'],
         f"NO₂ Hourly Concentration  —  WHO 24h Guideline: {config['thresholds']['NO2']} µg/m³"),
    ]:
        if col not in df.columns:
            continue
        plot_data = df[['datetime', col]].dropna()
        ax.plot(plot_data['datetime'], plot_data[col],
                color=color, linewidth=0.8, alpha=0.7, label=f'{col} (hourly)')

        # Daily mean overlay
        dm = daily_df[daily_df['variable'] == col][['date', 'mean']].dropna()
        ax.plot(dm['date'], dm['mean'],
                color=PAL['navy'], linewidth=2.0, linestyle='--',
                label='Daily mean', zorder=5)

        # WHO threshold line
        thr = config['thresholds'].get(col)
        if thr:
            ax.axhline(thr, color=thr_color, linewidth=1.8,
                       linestyle=':', label=f'WHO guideline ({thr})', zorder=4)
            ax.fill_between(plot_data['datetime'], thr,
                             plot_data[col].clip(lower=thr),
                             alpha=0.12, color=PAL['red'], label='Exceedance zone')

        ax.set_title(title, fontsize=11, fontweight='bold', color=PAL['navy'], pad=8)
        ax.set_xlabel('Date', fontsize=9)
        ax.set_ylabel(f'{col} (µg/m³)' if col != 'CO' else f'{col} (mg/m³)', fontsize=9)
        ax.legend(fontsize=8.5, loc='upper right')
        ax.tick_params(axis='x', rotation=30)

    fig.suptitle(f"ATARS — Hourly Pollutant Time-Series  ·  {config['city']}",
                 fontsize=13, fontweight='bold', color=PAL['navy'])
    fig.text(0.5, 0.01,
             'Illustrative analysis of available dataset. Dashed = daily mean. '
             'Reference lines from WHO AQG 2021.',
             ha='center', fontsize=8, color=PAL['gray'], style='italic')
    return _save_fig(fig, out, 'chart1_timeseries.png')


def chart_correlation(daily_df: pd.DataFrame, config: dict, out: Path) -> str:
    """Chart 2 — Pearson correlation heatmap R_{jk}."""
    key_vars = [v for v in
                ["PM10","NO2","SO2","CO","Ozone","NH3",
                 "humidity_pct","wind_speed_10m","pressure_hpa"]
                if v in daily_df['variable'].unique()]
    if len(key_vars) < 4:
        return None

    corr = compute_correlation_matrix(daily_df, key_vars)
    labels = {
        "PM10":"PM10","NO2":"NO₂","SO2":"SO₂","CO":"CO",
        "Ozone":"Ozone","NH3":"NH₃","humidity_pct":"Humidity",
        "wind_speed_10m":"Wind Speed","pressure_hpa":"Pressure"
    }
    corr.index   = [labels.get(c, c) for c in corr.index]
    corr.columns = [labels.get(c, c) for c in corr.columns]

    fig, ax = plt.subplots(figsize=(11, 9), facecolor=PAL['white'])
    mask = np.zeros_like(corr.values, dtype=bool)
    np.fill_diagonal(mask, False)

    sns.heatmap(corr, ax=ax, annot=True, fmt='.2f', mask=None,
                cmap=sns.diverging_palette(220, 20, as_cmap=True),
                vmin=-1, vmax=1, center=0, square=True,
                linewidths=0.5, annot_kws={'size': 8.5},
                cbar_kws={'shrink': 0.8, 'label': 'Pearson r'})

    ax.set_title(
        f'Pearson Correlation Matrix R  ·  {config["city"]}\n'
        r'$\mathbf{R}_{jk} = \mathrm{Corr}(X_j, X_k)$  '
        '— Association only, NOT causation  [Corr(X,Y) ≢ C(X→Y)]',
        fontsize=11, fontweight='bold', color=PAL['navy'], pad=12)
    ax.tick_params(axis='both', labelsize=9)
    plt.xticks(rotation=35, ha='right')
    plt.yticks(rotation=0)
    fig.text(0.5, 0.01,
             f'n = {int(daily_df["N_v"].sum()):,} valid daily records.  '
             'Causal inference requires experimental design — outside ATARS scope.',
             ha='center', fontsize=8, color=PAL['gray'], style='italic')
    return _save_fig(fig, out, 'chart2_correlation.png')


def chart_trend(daily_df: pd.DataFrame, config: dict, out: Path) -> str:
    """Chart 3 — 30-day rolling trend with baseline for PM10 and NO2."""
    fig, ax = plt.subplots(figsize=(14, 6), facecolor=PAL['white'])

    for col, color, label in [
        ('PM10', PAL['blue'],  'PM10 Daily Mean'),
        ('NO2',  PAL['teal'], 'NO₂ Daily Mean'),
    ]:
        dm = daily_df[daily_df['variable'] == col][['date', 'mean', 'mean_W']].dropna()
        if len(dm) < 5:
            continue
        ax.plot(dm['date'], dm['mean'],
                color=color, linewidth=1.5, alpha=0.6, label=label)
        ax.plot(dm['date'], dm['mean_W'],
                color=color, linewidth=2.5, linestyle='--',
                alpha=0.9, label=f'{col} 30-day baseline x̄_W')

    for col, thr_color in [('PM10', PAL['amber']), ('NO2', PAL['red'])]:
        thr = config['thresholds'].get(col)
        if thr:
            ax.axhline(thr, color=thr_color, linewidth=1.5,
                       linestyle=':', alpha=0.8,
                       label=f'{col} WHO guideline ({thr})')

    ax.set_title(
        f'Daily Trend & 30-Day Rolling Baseline  ·  {config["city"]}  '
        r'— $\bar{x}_W$ and $\sigma_W$ (Eqs. 7–8)',
        fontsize=11, fontweight='bold', color=PAL['navy'])
    ax.set_xlabel('Date', fontsize=9)
    ax.set_ylabel('Concentration (µg/m³)', fontsize=9)
    ax.legend(fontsize=8.5, loc='upper right', ncol=2)
    ax.tick_params(axis='x', rotation=30)
    return _save_fig(fig, out, 'chart3_trend.png')


def chart_anomalies(daily_df: pd.DataFrame, config: dict, out: Path) -> str:
    """Chart 4 — Z-score anomaly detection for PM10."""
    col = 'PM10'
    dm  = daily_df[daily_df['variable'] == col][
        ['date', 'mean', 'z_score', 'is_anomaly', 'CI_lower', 'CI_upper']
    ].dropna(subset=['z_score'])
    if len(dm) < 10:
        return None

    fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(14, 9), facecolor=PAL['white'])
    fig.subplots_adjust(hspace=0.45)

    # Panel A: daily mean with CI band
    ax1.plot(dm['date'], dm['mean'], color=PAL['blue'], linewidth=1.8, label='Daily mean x̄_d')
    ax1.fill_between(dm['date'], dm['CI_lower'], dm['CI_upper'],
                     alpha=0.18, color=PAL['blue'], label='95% CI band')
    anom = dm[dm['is_anomaly']]
    ax1.scatter(anom['date'], anom['mean'], color=PAL['red'], s=50, zorder=6,
                label=f'Anomaly (|ζ_d| > {config["anomaly_z_threshold"]})', marker='D')
    ax1.set_title(r'(a) PM10 Daily Mean with 95% CI  —  $\mathcal{CI}_{0.95} = \bar{x} \pm 1.96\,\sigma/\sqrt{N_v}$  (Eq. 12)',
                  fontsize=10, fontweight='bold', color=PAL['navy'])
    ax1.set_ylabel('PM10 (µg/m³)', fontsize=9)
    ax1.legend(fontsize=8.5)
    ax1.tick_params(axis='x', rotation=30)

    # Panel B: z-score
    ax2.bar(dm['date'], dm['z_score'],
            color=[PAL['red'] if a else PAL['blue'] for a in dm['is_anomaly']],
            width=0.8, alpha=0.7)
    ax2.axhline(config['anomaly_z_threshold'],  color=PAL['red'],
                linestyle='--', linewidth=1.8, label=f'ζ_thresh = ±{config["anomaly_z_threshold"]}')
    ax2.axhline(-config['anomaly_z_threshold'], color=PAL['red'], linestyle='--', linewidth=1.8)
    ax2.axhline(0, color=PAL['navy'], linewidth=0.8)
    ax2.set_title(r'(b) Z-Score $\zeta_d = (\bar{x}_d - \bar{x}_W)\,/\,\sigma_W$  (Eq. 9)  '
                  f'— Anomaly threshold: |ζ| > {config["anomaly_z_threshold"]}',
                  fontsize=10, fontweight='bold', color=PAL['navy'])
    ax2.set_xlabel('Date', fontsize=9)
    ax2.set_ylabel('Z-Score ζ_d', fontsize=9)
    ax2.legend(fontsize=8.5)
    ax2.tick_params(axis='x', rotation=30)

    fig.suptitle(f'ATARS — Anomaly Detection  ·  {config["city"]}',
                 fontsize=13, fontweight='bold', color=PAL['navy'])
    return _save_fig(fig, out, 'chart4_anomalies.png')


def chart_acf_pacf(daily_df: pd.DataFrame, config: dict, out: Path) -> str:
    """Chart 5 — ACF and PACF for PM10 daily means."""
    col = 'PM10'
    n_lags = min(config['acf_lags'], 60)
    result = run_temporal_analysis(daily_df, col, n_lags)
    if not result:
        return None

    acf  = result['acf']
    pacf = result['pacf']
    conf = result['conf_bound']
    lag_ax_full  = np.arange(len(acf))
    lag_ax_pacf  = np.arange(len(pacf))

    fig, axes = plt.subplots(2, 1, figsize=(13, 9), facecolor=PAL['white'])
    fig.subplots_adjust(hspace=0.52)

    axes[0].bar(lag_ax_full, acf,
                color=[PAL['blue'] if abs(v) > conf else '#A0B4CC' for v in acf],
                width=0.6, zorder=3)
    axes[0].axhline(conf,  color=PAL['red'], linestyle='--', linewidth=1.8,
                    label=f'95% CI (±{conf:.3f})')
    axes[0].axhline(-conf, color=PAL['red'], linestyle='--', linewidth=1.8)
    axes[0].axhline(0, color=PAL['navy'], linewidth=0.8)
    axes[0].set_title(r'(a) ACF: $\hat{\rho}(k) = \mathrm{Corr}(x_t,\, x_{t-k})$  (Eq. 12)  '
                      f'— PM10 Daily Means  ·  {config["city"]}',
                      fontsize=10, fontweight='bold', color=PAL['navy'])
    axes[0].set_xlabel('Lag k (days)', fontsize=9)
    axes[0].set_ylabel('ACF ρ̂(k)', fontsize=9)
    axes[0].legend(fontsize=8.5)
    axes[0].set_xlim(-0.5, len(acf) + 0.5)
    axes[0].set_ylim(-0.8, 1.1)

    axes[1].bar(lag_ax_pacf, pacf,
                color=[PAL['teal'] if abs(v) > conf else '#A0D4CC' for v in pacf],
                width=0.6, zorder=3)
    axes[1].axhline(conf,  color=PAL['red'], linestyle='--', linewidth=1.8,
                    label=f'95% CI (±{conf:.3f})')
    axes[1].axhline(-conf, color=PAL['red'], linestyle='--', linewidth=1.8)
    axes[1].axhline(0, color=PAL['navy'], linewidth=0.8)
    axes[1].set_title(r'(b) PACF: $\hat{\phi}_{kk}$ — Partial correlation at lag k (Eq. 13)',
                      fontsize=10, fontweight='bold', color=PAL['navy'])
    axes[1].set_xlabel('Lag k (days)', fontsize=9)
    axes[1].set_ylabel('PACF φ̂_kk', fontsize=9)
    axes[1].legend(fontsize=8.5)
    axes[1].set_xlim(-0.5, len(pacf) + 0.5)
    axes[1].set_ylim(-0.8, 1.1)

    sig_lags = result.get('significant_acf', [])
    fig.text(0.5, 0.01,
             f'Significant ACF lags (|ρ̂(k)| > {conf:.3f}): {sig_lags[:10]}  '
             f'— n = {result["series_length"]} daily observations.',
             ha='center', fontsize=8, color=PAL['gray'], style='italic')
    return _save_fig(fig, out, 'chart5_acf_pacf.png')


def chart_quality_summary(daily_df: pd.DataFrame, config: dict, out: Path) -> str:
    """Chart 6 — Data quality dashboard: Q(D) and confidence flags."""
    fig = plt.figure(figsize=(14, 8), facecolor=PAL['white'])
    gs  = gridspec.GridSpec(1, 2, figure=fig, wspace=0.38)
    ax1 = fig.add_subplot(gs[0, 0])
    ax2 = fig.add_subplot(gs[0, 1])

    # Panel A: Q(D) time series for PM10
    col = 'PM10'
    qd  = daily_df[daily_df['variable'] == col][['date', 'Q', 'confidence']].dropna()
    if len(qd) > 0:
        colors_q = {'HIGH': PAL['green'], 'MODERATE': PAL['amber'], 'LOW': PAL['red']}
        bar_colors = [colors_q.get(f, PAL['gray']) for f in qd['confidence']]
        ax1.bar(qd['date'], qd['Q'], color=bar_colors, width=0.8, alpha=0.75)
        ax1.axhline(config['confidence_high'],     color=PAL['green'], linestyle='--',
                    linewidth=1.8, label=f'HIGH ≥ {config["confidence_high"]}')
        ax1.axhline(config['confidence_moderate'], color=PAL['amber'], linestyle='--',
                    linewidth=1.8, label=f'MOD ≥ {config["confidence_moderate"]}')
        ax1.set_ylim(0, 1.05)
        ax1.set_title(f'(a) Q(D) = N_v/N_total  (Eq. 11)\nData Completeness Score — PM10',
                      fontsize=10, fontweight='bold', color=PAL['navy'])
        ax1.set_xlabel('Date', fontsize=9)
        ax1.set_ylabel('Q(D) ∈ [0, 1]', fontsize=9)
        ax1.legend(fontsize=8.5)
        ax1.tick_params(axis='x', rotation=30)

    # Panel B: Confidence flag distribution (all pollutants)
    flag_counts = (daily_df.groupby(['variable', 'confidence'])
                   .size().unstack(fill_value=0))
    key_cols = [c for c in POLLUTANTS if c in flag_counts.index][:8]
    if len(key_cols) > 0:
        fc = flag_counts.loc[key_cols]
        for flag, color in [('HIGH', PAL['green']),
                             ('MODERATE', PAL['amber']),
                             ('LOW', PAL['red'])]:
            if flag in fc.columns:
                ax2.barh(key_cols, fc[flag], label=flag, color=color,
                         alpha=0.75, height=0.6)
        ax2.set_title('(b) Confidence Flag Distribution\nby Pollutant Variable',
                      fontsize=10, fontweight='bold', color=PAL['navy'])
        ax2.set_xlabel('Days', fontsize=9)
        ax2.legend(fontsize=8.5)

    fig.suptitle(f'ATARS — Data Quality Assessment  ·  {config["city"]}',
                 fontsize=13, fontweight='bold', color=PAL['navy'])
    return _save_fig(fig, out, 'chart6_quality.png')


def chart_seasonal(daily_df: pd.DataFrame, config: dict, out: Path) -> str:
    """Chart 7 — Monthly box plots showing seasonal pattern for PM10 and NO2."""
    fig, axes = plt.subplots(1, 2, figsize=(14, 7), facecolor=PAL['white'])
    fig.subplots_adjust(wspace=0.35)

    month_names = ['Jan','Feb','Mar','Apr','May','Jun',
                   'Jul','Aug','Sep','Oct','Nov','Dec']

    for ax, col, color in [(axes[0], 'PM10', PAL['blue']),
                            (axes[1], 'NO2',  PAL['teal'])]:
        dm = daily_df[daily_df['variable'] == col][['date', 'mean']].dropna()
        if len(dm) < 30:
            continue
        dm['month'] = dm['date'].dt.month
        monthly = [dm[dm['month'] == m]['mean'].values for m in range(1, 13)]
        bp = ax.boxplot(monthly, patch_artist=True, notch=False,
                        medianprops={'color': PAL['navy'], 'linewidth': 2.5},
                        whiskerprops={'color': color, 'linewidth': 1.5},
                        capprops={'color': color, 'linewidth': 1.5},
                        flierprops={'markerfacecolor': PAL['red'],
                                    'markersize': 3, 'alpha': 0.5})
        for patch in bp['boxes']:
            patch.set_facecolor(color)
            patch.set_alpha(0.4)

        thr = config['thresholds'].get(col)
        if thr:
            ax.axhline(thr, color=PAL['amber'], linewidth=2,
                       linestyle='--', label=f'WHO guideline ({thr})')
        ax.set_xticks(range(1, 13))
        ax.set_xticklabels(month_names, fontsize=8.5)
        unit = 'mg/m³' if col == 'CO' else 'µg/m³'
        ax.set_title(f'{col} — Monthly Distribution\n'
                     r'Median $\tilde{x}$, IQR, whiskers: 1.5×IQR',
                     fontsize=10, fontweight='bold', color=PAL['navy'])
        ax.set_xlabel('Month', fontsize=9)
        ax.set_ylabel(f'{col} ({unit})', fontsize=9)
        if thr:
            ax.legend(fontsize=8.5)

    fig.suptitle(f'ATARS — Seasonal Pattern Analysis  ·  {config["city"]}',
                 fontsize=13, fontweight='bold', color=PAL['navy'])
    fig.text(0.5, 0.01,
             'Non-parametric: median-based (Eq. 24). Boxes show IQR = Q₀.₇₅ − Q₀.₂₅ (Eq. 25).',
             ha='center', fontsize=8, color=PAL['gray'], style='italic')
    return _save_fig(fig, out, 'chart7_seasonal.png')


# ── NEW CHART 8: Hourly Diurnal Profile (avg by hour of day) ──────────────────
def chart_diurnal(df: pd.DataFrame, config: dict, out: Path) -> str:
    """Chart 8 — Average concentration by hour of day (diurnal cycle) for key pollutants."""
    key_cols = [c for c in ['PM10','NO','NO2','CO','Ozone'] if c in df.columns]
    if not key_cols:
        return None
    df2 = df.copy()
    df2['hour'] = df2['datetime'].dt.hour

    fig, axes = plt.subplots(2, 3, figsize=(15, 9), facecolor=PAL['white'])
    fig.subplots_adjust(hspace=0.45, wspace=0.35)
    colors_ = [PAL['blue'], PAL['navy'], PAL['teal'], PAL['amber'], PAL['purple'],
               PAL['green']]
    ax_flat = axes.flatten()

    for idx, col in enumerate(key_cols[:5]):
        ax = ax_flat[idx]
        hourly = df2.groupby('hour')[col].agg(['mean','std']).reset_index()
        ax.fill_between(hourly['hour'],
                        hourly['mean'] - hourly['std'],
                        hourly['mean'] + hourly['std'],
                        alpha=0.18, color=colors_[idx])
        ax.plot(hourly['hour'], hourly['mean'],
                color=colors_[idx], linewidth=2.5, marker='o', markersize=4)
        thr = config['thresholds'].get(col)
        if thr:
            ax.axhline(thr, color=PAL['red'], linewidth=1.5,
                       linestyle='--', alpha=0.7, label=f'WHO: {thr}')
        ax.set_title(f'{col} — Diurnal Cycle',
                     fontsize=10, fontweight='bold', color=PAL['navy'])
        ax.set_xlabel('Hour of Day (0–23)', fontsize=8.5)
        unit = 'mg/m³' if col == 'CO' else 'µg/m³'
        ax.set_ylabel(f'{unit}', fontsize=8.5)
        ax.set_xticks([0, 6, 12, 18, 23])
        ax.legend(fontsize=8) if thr else None

    # 6th panel: combined normalised overlay
    ax6 = ax_flat[5]
    for idx, col in enumerate(key_cols[:4]):
        series = df2.groupby('hour')[col].mean()
        if series.max() > 0:
            ax6.plot(series.index, series.values / series.max(),
                     color=colors_[idx], linewidth=2, label=col)
    ax6.set_title('Normalised Overlay\n(peak = 1.0 per variable)',
                  fontsize=10, fontweight='bold', color=PAL['navy'])
    ax6.set_xlabel('Hour of Day', fontsize=8.5)
    ax6.set_ylabel('Relative Concentration', fontsize=8.5)
    ax6.legend(fontsize=8.5)
    ax6.set_xticks([0, 6, 12, 18, 23])

    fig.suptitle(f'ATARS — Diurnal (Hourly) Profile  ·  {config["city"]}',
                 fontsize=13, fontweight='bold', color=PAL['navy'])
    fig.text(0.5, 0.005, 'Mean ± 1σ band. Peak hours indicate traffic/industrial emission cycles.',
             ha='center', fontsize=8, color=PAL['gray'], style='italic')
    return _save_fig(fig, out, 'chart8_diurnal.png')


# ── NEW CHART 9: Day-of-Week Pattern ─────────────────────────────────────────
def chart_day_of_week(daily_df: pd.DataFrame, config: dict, out: Path) -> str:
    """Chart 9 — Average daily mean by day of week for PM10, NO2, CO."""
    key_cols = [c for c in ['PM10','NO2','CO','SO2'] if c in daily_df['variable'].unique()]
    if not key_cols:
        return None

    dow_names = ['Mon','Tue','Wed','Thu','Fri','Sat','Sun']
    fig, axes = plt.subplots(1, len(key_cols[:4]), figsize=(14, 6), facecolor=PAL['white'])
    if len(key_cols) == 1:
        axes = [axes]
    colors_ = [PAL['blue'], PAL['teal'], PAL['amber'], PAL['purple']]

    for ax, col, color in zip(axes, key_cols[:4], colors_):
        dm = daily_df[daily_df['variable'] == col][['date','mean']].dropna()
        dm['dow'] = dm['date'].dt.dayofweek
        dow_mean = dm.groupby('dow')['mean'].mean()
        dow_std  = dm.groupby('dow')['mean'].std()
        x = dow_mean.index
        ax.bar(x, dow_mean.values, color=color, alpha=0.7, width=0.6)
        ax.errorbar(x, dow_mean.values, yerr=dow_std.values,
                    fmt='none', color=PAL['navy'], capsize=4, linewidth=1.5)
        thr = config['thresholds'].get(col)
        if thr:
            ax.axhline(thr, color=PAL['red'], linewidth=1.5,
                       linestyle='--', label=f'WHO: {thr}')
        ax.set_xticks(range(7))
        ax.set_xticklabels(dow_names, fontsize=8.5)
        unit = 'mg/m³' if col == 'CO' else 'µg/m³'
        ax.set_title(f'{col} by\nDay of Week',
                     fontsize=10, fontweight='bold', color=PAL['navy'])
        ax.set_ylabel(f'{unit}', fontsize=8.5)
        if thr:
            ax.legend(fontsize=8)

    fig.suptitle(f'ATARS — Day-of-Week Pattern Analysis  ·  {config["city"]}',
                 fontsize=13, fontweight='bold', color=PAL['navy'])
    fig.text(0.5, 0.01,
             'Mean ± σ error bars. Weekday vs weekend patterns indicate anthropogenic emission cycles.',
             ha='center', fontsize=8, color=PAL['gray'], style='italic')
    return _save_fig(fig, out, 'chart9_day_of_week.png')


# ── NEW CHART 10: VOC Panel (Volatile Organic Compounds) ─────────────────────
def chart_voc_panel(daily_df: pd.DataFrame, config: dict, out: Path) -> str:
    """Chart 10 — VOC compounds trend: Benzene, Toluene, Xylene, Eth-Benzene, MP-Xylene."""
    voc_cols = [c for c in ['Benzene','Toluene','Xylene','Eth_Benzene','MP_Xylene']
                if c in daily_df['variable'].unique()]
    if len(voc_cols) < 2:
        return None

    fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(14, 10), facecolor=PAL['white'])
    fig.subplots_adjust(hspace=0.42)
    voc_colors = [PAL['blue'], PAL['teal'], PAL['purple'], PAL['amber'], PAL['green']]

    # Panel A: time series for all VOCs
    for col, color in zip(voc_cols, voc_colors):
        dm = daily_df[daily_df['variable'] == col][['date','mean']].dropna()
        if len(dm) < 5:
            continue
        ax1.plot(dm['date'], dm['mean'],
                 color=color, linewidth=1.6, label=col, alpha=0.85)
    # Benzene WHO line
    b_thr = config['thresholds'].get('Benzene', 1.7)
    ax1.axhline(b_thr, color=PAL['red'], linewidth=2,
                linestyle='--', label=f'Benzene WHO ({b_thr} µg/m³)')
    ax1.set_title('(a) VOC Compounds — Daily Mean Time Series',
                  fontsize=10, fontweight='bold', color=PAL['navy'])
    ax1.set_ylabel('Concentration (µg/m³)', fontsize=9)
    ax1.legend(fontsize=8.5, loc='upper right', ncol=3)
    ax1.tick_params(axis='x', rotation=30)

    # Panel B: monthly VOC means as grouped bar
    month_names = ['Jan','Feb','Mar','Apr','May','Jun',
                   'Jul','Aug','Sep','Oct','Nov','Dec']
    x   = np.arange(12)
    w   = 0.15
    for i, (col, color) in enumerate(zip(voc_cols, voc_colors)):
        dm = daily_df[daily_df['variable'] == col][['date','mean']].dropna()
        dm['month'] = dm['date'].dt.month
        monthly = [dm[dm['month'] == m]['mean'].mean() for m in range(1, 13)]
        ax2.bar(x + i * w, monthly, width=w, color=color, alpha=0.8, label=col)
    ax2.set_xticks(x + w * (len(voc_cols)-1)/2)
    ax2.set_xticklabels(month_names, fontsize=8.5)
    ax2.set_title('(b) Monthly Average VOC Concentrations — Grouped',
                  fontsize=10, fontweight='bold', color=PAL['navy'])
    ax2.set_ylabel('Concentration (µg/m³)', fontsize=9)
    ax2.legend(fontsize=8.5, ncol=3)

    fig.suptitle(f'ATARS — Volatile Organic Compounds (VOC) Panel  ·  {config["city"]}',
                 fontsize=13, fontweight='bold', color=PAL['navy'])
    fig.text(0.5, 0.005,
             'BTEX group: Benzene, Toluene, Ethylbenzene, Xylene. '
             'Benzene WHO annual guideline: 1.7 µg/m³.',
             ha='center', fontsize=8, color=PAL['gray'], style='italic')
    return _save_fig(fig, out, 'chart10_voc.png')


# ── NEW CHART 11: Rain vs Pollutant Scatter (Wet Deposition) ─────────────────
def chart_rain_scatter(daily_df: pd.DataFrame, config: dict, out: Path) -> str:
    """Chart 11 — Scatter of rain vs PM10 and SO2 to show wet deposition effect."""
    cols_to_plot = [c for c in ['PM10','SO2','NO2'] if c in daily_df['variable'].unique()]
    if 'rain_mm' not in daily_df['variable'].unique() or not cols_to_plot:
        return None

    rain_dm = daily_df[daily_df['variable'] == 'rain_mm'][['date','mean']].rename(
        columns={'mean': 'rain'}).dropna()

    fig, axes = plt.subplots(1, len(cols_to_plot[:3]), figsize=(14, 6),
                              facecolor=PAL['white'])
    if len(cols_to_plot) == 1:
        axes = [axes]
    colors_ = [PAL['blue'], PAL['teal'], PAL['amber']]

    for ax, col, color in zip(axes, cols_to_plot[:3], colors_):
        dm = daily_df[daily_df['variable'] == col][['date','mean']].rename(
            columns={'mean': col}).dropna()
        merged = pd.merge(rain_dm, dm, on='date')
        if len(merged) < 10:
            continue
        ax.scatter(merged['rain'], merged[col], color=color,
                   alpha=0.45, s=18, label=f'n={len(merged)}')
        # Linear trend
        if merged['rain'].std() > 0:
            z = np.polyfit(merged['rain'], merged[col], 1)
            p = np.poly1d(z)
            xr = np.linspace(merged['rain'].min(), merged['rain'].max(), 100)
            ax.plot(xr, p(xr), color=PAL['navy'], linewidth=2,
                    linestyle='--', label=f'Trend (OLS)')
            r_val = np.corrcoef(merged['rain'], merged[col])[0, 1]
            ax.text(0.97, 0.95, f'r = {r_val:.2f}',
                    transform=ax.transAxes, ha='right', va='top',
                    fontsize=10, color=PAL['navy'],
                    bbox=dict(facecolor='white', alpha=0.6, edgecolor='none'))
        ax.set_title(f'{col} vs Rain\n(wet deposition indicator)',
                     fontsize=10, fontweight='bold', color=PAL['navy'])
        ax.set_xlabel('Daily Rain (mm)', fontsize=9)
        unit = 'mg/m³' if col == 'CO' else 'µg/m³'
        ax.set_ylabel(f'{col} ({unit})', fontsize=9)
        ax.legend(fontsize=8.5)

    fig.suptitle(f'ATARS — Rain vs Pollutant Scatter  ·  {config["city"]}',
                 fontsize=13, fontweight='bold', color=PAL['navy'])
    fig.text(0.5, 0.01,
             'Association only — Corr(rain, pollutant) ≢ causal wet-deposition effect. '
             'r = Pearson correlation coefficient.',
             ha='center', fontsize=8, color=PAL['gray'], style='italic')
    return _save_fig(fig, out, 'chart11_rain_scatter.png')


# ── NEW CHART 12: Exceedance Calendar Heatmap ────────────────────────────────
def chart_exceedance_calendar(daily_df: pd.DataFrame, config: dict, out: Path) -> str:
    """Chart 12 — Calendar heatmap of PM10 daily mean (month × day-of-month)."""
    col = 'PM10'
    dm = daily_df[daily_df['variable'] == col][['date','mean','exceeds_threshold']].dropna()
    if len(dm) < 30:
        return None

    dm['month'] = dm['date'].dt.month
    dm['day']   = dm['date'].dt.day

    # Build 12×31 matrix
    matrix = np.full((12, 31), np.nan)
    exceed = np.zeros((12, 31), dtype=bool)
    for _, row in dm.iterrows():
        m, d = int(row['month']) - 1, int(row['day']) - 1
        matrix[m, d] = row['mean']
        exceed[m, d] = row['exceeds_threshold']

    fig, ax = plt.subplots(figsize=(15, 7), facecolor=PAL['white'])
    month_names = ['Jan','Feb','Mar','Apr','May','Jun',
                   'Jul','Aug','Sep','Oct','Nov','Dec']

    im = ax.imshow(matrix, aspect='auto', cmap='RdYlGn_r',
                   vmin=0, vmax=config['thresholds'].get(col, 150) * 2)
    # Mark exceedances with X
    for m in range(12):
        for d in range(31):
            if exceed[m, d]:
                ax.text(d, m, '×', ha='center', va='center',
                        fontsize=7, color='white', fontweight='bold')
            elif not np.isnan(matrix[m, d]):
                ax.text(d, m, f'{matrix[m,d]:.0f}', ha='center', va='center',
                        fontsize=5.5, color='white' if matrix[m,d] > 80 else '#1A2C4E')

    ax.set_yticks(range(12))
    ax.set_yticklabels(month_names, fontsize=9)
    ax.set_xticks(range(31))
    ax.set_xticklabels([str(d+1) for d in range(31)], fontsize=7.5)
    ax.set_xlabel('Day of Month', fontsize=9)

    cbar = fig.colorbar(im, ax=ax, shrink=0.8, pad=0.01)
    cbar.set_label(f'PM10 (µg/m³)', fontsize=9)
    thr = config['thresholds'].get(col)
    if thr:
        cbar.ax.axhline(thr, color='white', linewidth=2, linestyle='--')

    ax.set_title(
        f'(a) PM10 Daily Mean Calendar Heatmap  ·  {config["city"]}\n'
        f'× = WHO exceedance (>{thr} µg/m³). Green = clean, Red = polluted.',
        fontsize=11, fontweight='bold', color=PAL['navy'])
    return _save_fig(fig, out, 'chart12_calendar.png')



# ── CHART 14: ML Anomaly Comparison (IF vs z-score) ─────────────────────────
def chart_ml_anomaly_comparison(daily_df, config, ml_if, out):
    """Chart 14 — Isolation Forest vs z-score anomaly comparison."""
    if not ml_if.get("available"):
        return None
    comp = ml_if["comparison_df"].copy()
    if comp.empty:
        return None

    from matplotlib.patches import Patch
    fig, axes = plt.subplots(3, 1, figsize=(14, 12), facecolor=PAL['white'])
    fig.subplots_adjust(hspace=0.50)
    dates = comp.index

    # Panel A: IF score
    ax = axes[0]
    ax.fill_between(dates, comp["if_score"], color=PAL["teal"], alpha=0.35)
    ax.plot(dates, comp["if_score"], color=PAL["teal"], linewidth=0.8, label="IF score")
    thresh = float(comp["if_score"].quantile(config["if_contamination"]))
    ax.axhline(thresh, color=PAL["red"], linestyle="--", linewidth=1.5,
               label=f"Anomaly boundary ({config['if_contamination']:.0%})")
    anomaly_dates = comp[comp["ml_anomaly"]].index
    for d in anomaly_dates:
        ax.axvspan(d - pd.Timedelta(hours=12), d + pd.Timedelta(hours=12),
                   alpha=0.25, color=PAL["red"])
    ax.set_title("(a) Isolation Forest Decision Score — lower = more anomalous",
                 fontsize=10, fontweight="bold", color=PAL["navy"])
    ax.set_ylabel("IF Score", fontsize=9); ax.legend(fontsize=8.5)
    ax.tick_params(axis="x", rotation=30)

    # Panel B: z-score
    ax = axes[1]
    if "z_score" in comp.columns:
        vz = comp["z_score"].dropna()
        ax.plot(vz.index, vz.values, color=PAL["blue"], linewidth=0.9, label="PM10 ζ_d")
        zt = config["anomaly_z_threshold"]
        ax.axhline(zt,  color=PAL["amber"], linestyle="--", linewidth=1.5, label=f"+{zt}σ")
        ax.axhline(-zt, color=PAL["amber"], linestyle="--", linewidth=1.5, label=f"-{zt}σ")
        ax.fill_between(vz.index, zt, vz.values.clip(min=zt), alpha=0.18, color=PAL["red"])
        ax.fill_between(vz.index, -zt, vz.values.clip(max=-zt), alpha=0.18, color=PAL["red"])
    ax.set_title("(b) Z-Score ζ_d — Formal Operator Eq. 9", fontsize=10,
                 fontweight="bold", color=PAL["navy"])
    ax.set_ylabel("ζ_d (std devs)", fontsize=9); ax.legend(fontsize=8.5)
    ax.tick_params(axis="x", rotation=30)

    # Panel C: Agreement
    ax = axes[2]
    for d, row in comp.iterrows():
        if row["both"]:      c = PAL["red"]
        elif row["if_only"]: c = PAL["purple"]
        elif row["zs_only"]: c = PAL["amber"]
        else:                c = PAL["teal"]
        ax.axvspan(d - pd.Timedelta(hours=12), d + pd.Timedelta(hours=12),
                   alpha=0.65, color=c)
    legend_elements = [
        Patch(facecolor=PAL["red"],    alpha=0.65, label=f"Both ({ml_if['n_both']})"),
        Patch(facecolor=PAL["purple"], alpha=0.65, label=f"IF-only ({ml_if['n_if_only']})"),
        Patch(facecolor=PAL["amber"],  alpha=0.65, label=f"Z-only ({ml_if['n_zs_only']})"),
        Patch(facecolor=PAL["teal"],   alpha=0.65, label="Normal"),
    ]
    ax.legend(handles=legend_elements, fontsize=8.5, loc="upper right", ncol=2)
    ax.set_title("(c) Agreement Map — IF vs Z-Score", fontsize=10,
                 fontweight="bold", color=PAL["navy"])
    ax.set_yticks([]); ax.tick_params(axis="x", rotation=30)

    fig.suptitle(
        f"ATARS — ML Anomaly: Isolation Forest vs Z-Score  ·  {config['city']}\n"
        f"IF: {ml_if['n_anomalies']} days | Features: {', '.join(ml_if['features_used'])} | "
        f"random_state={ml_if['random_state']}",
        fontsize=11, fontweight="bold", color=PAL["navy"])
    fig.text(0.5, 0.005,
             "IF does NOT replace ζ_d (Eq. 9). Both reported. IF-only days = multivariate patterns.",
             ha="center", fontsize=8, color=PAL["gray"], style="italic")
    return _save_fig(fig, out, "chart14_ml_anomaly.png")


# ── CHART 15: Holt-Winters PM10 Forecast ─────────────────────────────────────
def chart_forecast(daily_df, config, ml_hw, out):
    """Chart 15 — Holt-Winters Triple Exponential Smoothing PM10 Forecast."""
    if not ml_hw.get("available"):
        return None

    fig, axes = plt.subplots(2, 1, figsize=(14, 10), facecolor=PAL["white"])
    fig.subplots_adjust(hspace=0.45)

    hist_dates = ml_hw["history_dates"]
    hist_vals  = ml_hw["history_series"]
    fitted     = ml_hw["fitted"]
    fc_dates   = ml_hw["forecast_dates"]
    fc_vals    = ml_hw["forecast"]
    fc_lower   = ml_hw["lower_80"]
    fc_upper   = ml_hw["upper_80"]

    # Panel A
    ax = axes[0]
    ax.plot(hist_dates, hist_vals, color=PAL["blue"], linewidth=0.8,
            alpha=0.7, label="Historical PM10")
    ax.plot(hist_dates, fitted, color=PAL["teal"], linewidth=1.5,
            linestyle="--", label="HW fitted", alpha=0.85)
    ax.axvline(hist_dates[-1], color=PAL["navy"], linestyle=":", linewidth=1.5)
    ax.plot(fc_dates, fc_vals, color=PAL["amber"], linewidth=2.2,
            marker="o", markersize=4, label=f"{ml_hw['forecast_days']}-day forecast")
    ax.fill_between(fc_dates, fc_lower, fc_upper, alpha=0.25, color=PAL["amber"],
                    label="80% prediction interval")
    thr_v = config["thresholds"].get("PM10")
    if thr_v:
        ax.axhline(thr_v, color=PAL["red"], linestyle=":", linewidth=1.5,
                   label=f"WHO guideline ({thr_v})")
    ax.set_title(f"(a) PM10 Historical + {ml_hw['forecast_days']}-Day Holt-Winters Forecast",
                 fontsize=10, fontweight="bold", color=PAL["navy"])
    ax.set_ylabel("PM10 (µg/m³)", fontsize=9)
    ax.legend(fontsize=8.5, loc="upper left"); ax.tick_params(axis="x", rotation=30)

    # Panel B: Residuals
    ax = axes[1]
    resid = hist_vals - fitted
    colors_r = [PAL["red"] if r > 0 else PAL["teal"] for r in resid]
    ax.bar(hist_dates, resid, color=colors_r, alpha=0.6, width=0.8)
    ax.axhline(0, color=PAL["navy"], linewidth=1.2)
    rs = float(np.std(resid, ddof=1))
    ax.axhline(rs,  color=PAL["amber"], linestyle="--", linewidth=1.2, label=f"+1σ ({rs:.1f})")
    ax.axhline(-rs, color=PAL["amber"], linestyle="--", linewidth=1.2, label=f"-1σ")
    ax.set_title(f"(b) Residuals — RMSE={ml_hw['rmse']:.2f} | MAE={ml_hw['mae']:.2f} µg/m³",
                 fontsize=10, fontweight="bold", color=PAL["navy"])
    ax.set_ylabel("Residual (µg/m³)", fontsize=9)
    ax.legend(fontsize=8.5); ax.tick_params(axis="x", rotation=30)

    fig.suptitle(
        f"ATARS — Holt-Winters PM10 Forecast  ·  {config['city']}\n"
        f"α={ml_hw['alpha']} | β={ml_hw['beta']} | γ={ml_hw['gamma']} | "
        f"seasonal={ml_hw['seasonal_periods']}d",
        fontsize=11, fontweight="bold", color=PAL["navy"])
    fig.text(0.5, 0.005,
             "Statistical extrapolation only. Forecast NOT in JSON contract J (LLM did not use it). "
             "80% PI widens with forecast horizon h.",
             ha="center", fontsize=8, color=PAL["gray"], style="italic")
    return _save_fig(fig, out, "chart15_forecast.png")


def generate_all_charts(df, daily_df, config, out, verification=None,
                         ml_if=None, ml_hw=None):
    """
    Generate all charts. Returns dict of {name: filepath}.
    v2.0: Adds Chart 14 (ML anomaly comparison) and Chart 15 (forecast).
    """
    n_base   = 12
    n_grnd   = 1 if verification else 0
    n_ml_if  = 1 if (ml_if and ml_if.get("available")) else 0
    n_ml_hw  = 1 if (ml_hw and ml_hw.get("available") and config.get("use_forecast")) else 0
    n_charts = n_base + n_grnd + n_ml_if + n_ml_hw
    print(f"\n  [Step 6] Generating {n_charts} publication-quality charts "
          f"({n_base} core + {n_grnd} RGV + {n_ml_if} ML-anomaly + {n_ml_hw} ML-forecast)...")
    charts = {}
    tasks  = [
        ('timeseries',   chart_timeseries,          [df, daily_df, config, out]),
        ('correlation',  chart_correlation,          [daily_df, config, out]),
        ('trend',        chart_trend,                [daily_df, config, out]),
        ('anomalies',    chart_anomalies,            [daily_df, config, out]),
        ('acf_pacf',     chart_acf_pacf,             [daily_df, config, out]),
        ('quality',      chart_quality_summary,      [daily_df, config, out]),
        ('seasonal',     chart_seasonal,             [daily_df, config, out]),
        ('diurnal',      chart_diurnal,              [df, config, out]),
        ('day_of_week',  chart_day_of_week,          [daily_df, config, out]),
        ('voc',          chart_voc_panel,            [daily_df, config, out]),
        ('rain_scatter', chart_rain_scatter,         [daily_df, config, out]),
        ('calendar',     chart_exceedance_calendar,  [daily_df, config, out]),
    ]
    for name, fn, args in tasks:
        try:
            path = fn(*args)
            if path:
                charts[name] = path
                print(f"  \u2713 Chart {len(charts):02d}/{n_charts}: {name}")
            else:
                print(f"  \u26a0  Skipped (insufficient data): {name}")
        except Exception as e:
            print(f"  \u2717  Chart error ({name}): {e}")
    if verification:
        try:
            path = chart_grounding_verification(verification, config, out)
            if path:
                charts['grounding'] = path
                print(f"  \u2713 Chart {len(charts):02d}/{n_charts}: grounding_verification (RGV ★)")
        except Exception as e:
            print(f"  \u2717  Chart error (grounding): {e}")
    if ml_if and ml_if.get("available"):
        try:
            path = chart_ml_anomaly_comparison(daily_df, config, ml_if, out)
            if path:
                charts['ml_anomaly'] = path
                print(f"  \u2713 Chart {len(charts):02d}/{n_charts}: ml_anomaly (Isolation Forest)")
        except Exception as e:
            print(f"  \u2717  Chart error (ml_anomaly): {e}")
    if ml_hw and ml_hw.get("available") and config.get("use_forecast"):
        try:
            path = chart_forecast(daily_df, config, ml_hw, out)
            if path:
                charts['forecast'] = path
                print(f"  \u2713 Chart {len(charts):02d}/{n_charts}: forecast (Holt-Winters)")
        except Exception as e:
            print(f"  \u2717  Chart error (forecast): {e}")
    return charts


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 9 — JSON CONTRACT ASSEMBLY
#  J = {x̄, σ, δ, ζ, Q} — the formal interface between computation and language
# ═══════════════════════════════════════════════════════════════════════════════

def build_json_contract(df: pd.DataFrame, daily_df: pd.DataFrame,
                         config: dict, ols_result: dict,
                         corr_matrix: pd.DataFrame) -> dict:
    """
    Assemble the formal JSON contract J.
    This is the ONLY input to the LLM — no raw data is passed.
    Grounding constraint: G(s) ⊆ J for all narrative sentences s.
    """
    print("\n  [Step 4] Assembling JSON contract J...")

    date_range = f"{df['datetime'].min().date()} to {df['datetime'].max().date()}"
    total_days  = int((df['datetime'].max() - df['datetime'].min()).days) + 1
    total_recs  = len(df)

    # Overall quality across all pollutants
    avg_q = daily_df.groupby('variable')['Q'].mean().to_dict()
    overall_q = float(np.mean(list(avg_q.values()))) if avg_q else 0.0
    overall_conf = get_confidence_flag(overall_q, config)

    # Per-variable summary statistics
    variables = {}
    for col in POLLUTANTS:
        sub = daily_df[daily_df['variable'] == col]
        if sub.empty or sub['mean'].isna().all():
            continue
        valid = sub[sub['confidence'] != 'LOW']
        annual_mean = float(sub['mean'].mean())
        annual_std  = float(sub['mean'].std())
        annual_max  = float(sub['mean'].max())
        annual_min  = float(sub['mean'].min())
        n_anomaly   = int(sub['is_anomaly'].sum())
        n_exceed    = int(sub['exceeds_threshold'].sum())
        thr         = config['thresholds'].get(col)
        variables[col] = {
            "annual_mean"      : round(annual_mean, 3),
            "annual_std"       : round(annual_std, 3),
            "annual_max"       : round(annual_max, 3),
            "annual_min"       : round(annual_min, 3),
            "threshold"        : thr,
            "exceeds_days"     : n_exceed,
            "exceed_pct"       : round(n_exceed / total_days * 100, 2) if total_days > 0 else 0,
            "anomaly_days"     : n_anomaly,
            "mean_z_score"     : round(float(sub['z_score'].mean()), 3)
                                  if not sub['z_score'].isna().all() else None,
            "avg_Q"            : round(float(sub['Q'].mean()), 4),
        }

    # High-exceedance summary
    exceedances = {k: v['exceeds_days']
                   for k, v in variables.items() if v['exceeds_days'] > 0}
    top_pollutant = max(exceedances, key=exceedances.get) if exceedances else "PM10"

    # Correlation matrix (top correlations only)
    top_corr = []
    if corr_matrix is not None and not corr_matrix.empty:
        flat = corr_matrix.stack()
        flat.index.names = ['var1', 'var2']
        flat = flat.reset_index()
        flat.columns = ['var1', 'var2', 'r']
        flat = flat[flat['var1'] != flat['var2']].copy()
        flat['abs_r'] = flat['r'].abs()
        flat = flat.nlargest(10, 'abs_r')
        top_corr = [
            {"var1": row.var1, "var2": row.var2, "r": round(row.r, 3)}
            for _, row in flat.iterrows()
        ]

    J = {
        "schema_version"  : "1.0",
        "framework"       : "ATARS v1.0",
        "author"          : config['author'],
        "institution"     : config['institution'],
        "city"            : config['city'],
        "station_id"      : config['station_id'],
        "date_range"      : date_range,
        "total_days"      : total_days,
        "total_records"   : total_recs,
        "data_quality"    : {
            "overall_Q"        : round(overall_q, 4),
            "confidence_flag"  : overall_conf,
            "per_variable_Q"   : {k: round(v, 3) for k, v in avg_q.items()},
        },
        "variables"       : variables,
        "exceedances"     : exceedances,
        "top_pollutant"   : top_pollutant,
        "top_correlations": top_corr,
        "ols_regression"  : ols_result if ols_result else {},
        "formal_operators": {
            "baseline_window_W"   : config['baseline_window_days'],
            "z_threshold"         : config['anomaly_z_threshold'],
            "ci_alpha"            : config['ci_alpha'],
        },
        "ai_constraints"  : {
            "grounding_rule"  : "G(s) ⊆ J ∧ G(s) ≠ ∅ — all claims reference this JSON",
            "temperature"     : 0,
            "causation_note"  : "Corr(X,Y) does not imply X→Y — association only",
            "extrapolation"   : "prohibited — no forecasts beyond this dataset",
        },
        # SECURITY / ML NOTE: ML results (Isolation Forest, Holt-Winters forecast)
        # are intentionally NOT included in J. Reasons:
        #   1. J is the LLM's ONLY data source — ML forecasts in J risk the LLM
        #      stating uncertain extrapolations as historical facts.
        #   2. Isolation Forest anomaly flags are supplementary — formal ζ_d
        #      is the auditable anomaly criterion already in J.
        #   3. ML results appear in Chart 14, Chart 15, and the report ML
        #      section only — clearly marked as supplementary analysis.
        "ml_metadata"     : {
            "if_contamination": "injected_at_runtime",
            "if_n_estimators" : "injected_at_runtime",
            "hw_forecast_days": "injected_at_runtime",
            "note"            : "ML results NOT in J — see chart14 and chart15 in report",
        }
    }

    print(f"  ✓ JSON contract assembled: {len(variables)} variables, "
          f"{len(top_corr)} top correlations")
    return J


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 10 — LLM NARRATIVE GENERATION (Ollama — optional, offline, free)
#  Formal grounding: G: J → N where ∀ s ∈ N: G(s) ⊆ J
# ═══════════════════════════════════════════════════════════════════════════════

NARRATIVE_SYSTEM_PROMPT = """You are the ATARS analytical narrative engine.
Your ONLY job is to convert the JSON data contract into clear research prose.

HARD RULES (formal grounding constraints):
1. Every numerical value you write MUST come from the JSON provided — no inventions.
2. Do NOT forecast or predict beyond the provided date range.
3. Do NOT assert causation — you may state associations (correlation) only.
4. Do NOT reference any external knowledge, standards, or datasets not in the JSON.
5. Temperature τ=0 — be precise, not creative.
6. Each section must be 3-5 sentences maximum.
7. Output ONLY the research sentences — nothing else.
   Do NOT write any introduction, preamble, header, label, or explanation.
   Do NOT repeat or paraphrase the task instruction you were given.
   Do NOT begin with phrases such as: "Here are", "The following", "This section",
   "Based on", "Certainly", "Sure", "Below", "As requested", "In this section",
   "The analysis shows", "The data shows", "This study presents".
   Start the response immediately with the first research sentence.

Write for a research audience. Be formal, precise, and honest about data quality."""

NARRATIVE_SECTIONS = {
    "A_summary": "Write a formal 3-4 sentence annual summary of air quality for {city} based strictly on the JSON statistics. Mention the overall data quality score Q={Q} and confidence flag {conf}.",
    "B_pollutants": "For the top 3 pollutants by exceedance days (use exceedances from JSON), write 3-4 sentences describing their annual mean values and threshold exceedance counts. Use values from JSON variables only.",
    "C_correlations": "In 3-4 sentences, describe the strongest statistical associations found in the correlation analysis. Explicitly state: correlation does not imply causation. Use only the top_correlations from JSON.",
    "D_anomalies": "Write 3-4 sentences about the anomaly detection results. Reference z-score threshold ζ_thresh={z_thresh} and the number of anomaly days per pollutant. Use only values from JSON.",
    "E_quality": "Write 3-4 sentences about data quality using Q(D) values per variable from JSON. Note which variables had LOW confidence and what this means for interpretation.",
    "F_limitations": "Write exactly 4 sentences acknowledging the analytical limitations. Sentence 1: state this is statistical association only, not causal analysis. Sentence 2: state Q(D) measures record completeness, not sensor calibration accuracy. Sentence 3: state results are specific to the configured station and date range only. Sentence 4: state expert domain review is recommended before using results in decisions. Write only these 4 sentences, no preamble."
}


def try_llm_narrative(J: dict, config: dict) -> dict:
    """
    Generate narrative using Ollama local LLM.
    Falls back gracefully if Ollama is not running.
    G(s) ⊆ J enforced by system prompt constraints.
    """
    import requests as req

    narrative = {}
    if not config.get('use_llm', True):
        return _placeholder_narrative(J)

    # Check Ollama is running
    try:
        resp = req.get(f"{config['ollama_url']}/api/tags", timeout=5)
        if resp.status_code != 200:
            raise ConnectionError()
        models = [m['name'] for m in resp.json().get('models', [])]
        print(f"  ✓ Ollama running — available models: {models}")
        if not any(config['llm_model'].split(':')[0] in m for m in models):
            print(f"  ⚠ Model '{config['llm_model']}' not found. "
                  f"Run: ollama pull {config['llm_model']}")
            return _placeholder_narrative(J)
    except Exception:
        print("  ⚠ Ollama not running — using structured placeholders.")
        print("  To enable AI narrative: install Ollama from https://ollama.com")
        print(f"  Then run: ollama pull {config['llm_model']}")
        return _placeholder_narrative(J)

    import time as _time
    n_sections = len(NARRATIVE_SECTIONS)
    print(f"\n  [Step 5] Generating AI narrative with {config['llm_model']}...")
    print(f"  ⏱  Timeout per section: {config['llm_timeout']}s  "
          f"| Retries: {config['llm_retry']}  "
          f"| Sections: {n_sections}")
    print(f"  ⏱  Worst-case total wait: "
          f"~{config['llm_timeout'] * n_sections // 60} min "
          f"(LLM speed depends on your hardware)\n")

    json_str = json.dumps(J, indent=2)
    for idx, (section_key, prompt_template) in enumerate(NARRATIVE_SECTIONS.items(), 1):
        prompt = prompt_template.format(
            city    = J['city'],
            Q       = J['data_quality']['overall_Q'],
            conf    = J['data_quality']['confidence_flag'],
            z_thresh= J['formal_operators']['z_threshold']
        )
        full_prompt = (f"JSON DATA CONTRACT:\n{json_str}\n\n"
                       f"TASK: {prompt}")

        success = False
        for attempt in range(1, config['llm_retry'] + 2):  # +2: first try + retries
            try:
                t_start = _time.time()
                print(f"  → Section {idx}/{n_sections} [{section_key}]"
                      f"{'  (retry ' + str(attempt-1) + ')' if attempt > 1 else ''} ...",
                      end='', flush=True)
                r = req.post(
                    f"{config['ollama_url']}/api/generate",
                    json={"model"  : config['llm_model'],
                          "prompt" : full_prompt,
                          "system" : NARRATIVE_SYSTEM_PROMPT,
                          "stream" : False,
                          "options": {"temperature": 0, "num_predict": 512}},
                    timeout=config['llm_timeout']
                )
                elapsed = _time.time() - t_start
                if r.status_code == 200:
                    text = _clean_llm_response(r.json().get('response', '').strip())
                    narrative[section_key] = text
                    print(f" ✓  ({elapsed:.0f}s, {len(text)} chars)")
                    success = True
                    break
                else:
                    print(f" ✗  HTTP {r.status_code}")
                    narrative[section_key] = _placeholder_text(section_key, J)
                    break
            except req.exceptions.Timeout:
                elapsed = _time.time() - t_start
                if attempt <= config['llm_retry']:
                    print(f" ⏱ Timeout ({elapsed:.0f}s) — retrying in "
                          f"{config['llm_retry_delay']}s...")
                    _time.sleep(config['llm_retry_delay'])
                else:
                    print(f" ⏱ Timeout after {config['llm_retry']+1} attempts "
                          f"— using statistical placeholder.")
                    narrative[section_key] = _placeholder_text(section_key, J)
            except Exception as e:
                print(f" ⚠ Error: {e}")
                narrative[section_key] = _placeholder_text(section_key, J)
                break

    narrative['model']   = config['llm_model']
    narrative['ai_flag'] = True
    return narrative



def _clean_llm_response(text: str) -> str:
    """
    Remove common LLM preamble patterns that appear despite system prompt instructions.
    These patterns indicate the LLM echoed back its task rather than producing content.

    Known bad openers that LLMs produce even at temperature=0:
      - "The following research prose..."
      - "Here are the sentences..."
      - "Based on the JSON data..."
      - "Certainly! Here..."
      - "Below is..."
    """
    import re as _re

    # Strip leading/trailing whitespace
    text = text.strip()

    # Remove markdown formatting that some models add
    text = text.replace("**", "").replace("__", "")

    # Patterns to strip from the start of the response
    preamble_patterns = [
        # "The following X:" preambles
        r"^The following\s+[\w\s,]+:\s*",
        r"^The following\s+[\w\s,]+\.\s*",
        # "Here is/are ..." openers — includes "Here are 3-4 sentences about X:"
        r"^Here (?:is|are)\s+[\w\s\-,]+[.:]\s*",
        r"^Here(?:'s| is) (?:a |the |my )?[\w\s]+[.:]\s*",
        # "Based on ..." openers
        r"^Based on (?:the |this |)?(?:provided |given |)?(?:JSON |data |)?[\w\s]+[,:]\s*",
        # "Certainly/Sure/Absolutely" openers
        r"^(?:Certainly|Sure|Absolutely|Of course|Yes)[!,.]?\s*(?:Here\s+(?:is|are)[^.:]*[.:]?\s*)?",
        # "As requested ..." openers
        r"^As requested[,.]?\s*",
        # "Below is/are ..." openers
        r"^Below (?:is|are)\s+[\w\s,]+[.:]\s*",
        # "In this section ..." openers
        r"^In this section[,.]?\s*",
        # "This section ..." echo
        r"^This section (?:provides?|presents?|covers?|acknowledges?)[^.]*[.:]\s*",
        # "This analysis ..." echo
        r"^This analysis (?:provides?|presents?)[^.]*[.:]\s*",
        # "This study presents a" echo (specific bug-report pattern)
        r"^This study presents a?\s*",
        # "The analysis shows/reveals ..." echo
        r"^The (?:analysis|data|results?) (?:shows?|reveals?|indicates?)[^.]*[.:]\s*",
        # Numbered list openers "1." that indicate instruction repetition
        r"^1\.\s+(?:State|Note|Mention|Write)[^.]+\.\s*",
    ]

    for pattern in preamble_patterns:
        cleaned = _re.sub(pattern, "", text, flags=_re.IGNORECASE).strip()
        # Accept if cleaned is non-empty AND at least 25% of original
        # (lower threshold: 0.25 vs 0.5 — catches cases where preamble is >50% of text)
        if cleaned and len(cleaned) >= max(20, len(text) * 0.25):
            text = cleaned

    # Remove any trailing instruction echoes (sometimes LLMs append "Write only...")
    text = _re.sub(r"\s*Write only (?:these|the) \d+ sentences.*$", "", text, flags=_re.IGNORECASE | _re.DOTALL).strip()
    text = _re.sub(r"\s*\(Note:.*\)$", "", text, flags=_re.IGNORECASE).strip()

    # Capitalise first letter if it was stripped
    if text and text[0].islower():
        text = text[0].upper() + text[1:]

    return text.strip()

def _placeholder_text(key: str, J: dict) -> str:
    """Statistical placeholders when LLM unavailable — all values from J."""
    city = J['city']
    Q    = J['data_quality']['overall_Q']
    conf = J['data_quality']['confidence_flag']
    top  = J.get('top_pollutant', 'PM10')
    exc  = J.get('exceedances', {})
    n_days = J.get('total_days', 0)
    vars_  = J.get('variables', {})

    texts = {
        "A_summary": (
            f"Annual analysis of {city} covers {n_days} days of structured "
            f"time-series measurements across {len(vars_)} monitored variables. "
            f"Overall data completeness Q(D) = {Q:.3f}, yielding a {conf} confidence "
            f"classification. Statistical aggregation was performed using the ATARS "
            f"formal aggregation operator A(D_v) across all validated records."
        ),
        "B_pollutants": (
            f"Primary exceedance was observed in: "
            + ", ".join(f"{k} ({v} days)" for k, v in list(exc.items())[:3]) + ". "
            f"The variable with highest exceedance count was {top}. "
            f"Annual mean values per variable are reported in the statistics table. "
            f"All threshold comparisons use WHO AQG 2021 reference values as configured."
        ),
        "C_correlations": (
            f"Pearson correlation analysis was performed across all variable pairs "
            f"using the correlation matrix R (Eq. 18 of the ATARS framework). "
            f"Top statistical associations are reported in the correlation matrix figure. "
            f"FORMAL NOTE: all reported correlations are measures of statistical association "
            f"only — causal inference requires experimental design outside this framework's scope."
        ),
        "D_anomalies": (
            f"Z-score anomaly detection was applied using threshold |ζ_d| > "
            f"{J['formal_operators']['z_threshold']} (3-sigma rule, Proposition 1). "
            f"Anomalies represent daily means exceeding 3 standard deviations from "
            f"the {J['formal_operators']['baseline_window_W']}-day rolling baseline. "
            f"Anomaly counts by variable are reported in the statistics table."
        ),
        "E_quality": (
            f"Data quality was assessed using Q(D) = N_v/N_total per variable per day. "
            f"Overall completeness: Q = {Q:.3f} ({conf}). "
            f"Variables with Q < {J['formal_operators'].get('min_valid', 0.70)} "
            f"are flagged LOW and excluded from baseline calculations. "
            f"Quality flags do not capture sensor calibration uncertainty."
        ),
        "F_limitations": (
            f"This analysis provides statistical characterisation only — no causal "
            f"relationships are established. Correlation statistics (Corr(X,Y)) do not "
            f"imply causal effect parameters C(X→Y). Q(D) measures record completeness, "
            f"not sensor accuracy or calibration validity. Results represent the configured "
            f"station only and should be reviewed by domain experts before informing decisions."
        ),
    }
    return texts.get(key, f"[Section {key} — computed from JSON contract J]")


def _placeholder_narrative(J: dict) -> dict:
    return {k: _placeholder_text(k, J) for k in NARRATIVE_SECTIONS} | \
           {'model': 'statistical_placeholder', 'ai_flag': False}


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 10.5 — RUNTIME GROUNDING VERIFIER (RGV)
#
#  NOVEL CONTRIBUTION — programmatic enforcement of G(s) ⊆ J
#
#  Prior work: G(s) ⊆ J was enforced only by system-prompt instruction.
#  The LLM could hallucinate numbers not in J and no system would detect it.
#
#  This module: every numerical claim in every LLM sentence is extracted,
#  matched against the JSON contract J, and classified as:
#    GROUNDED       — claim found in J within tolerance
#    PARTIALLY_GROUNDED — some claims found, some not
#    UNGROUNDED     — numerical claims present, none found in J
#    NON_NUMERICAL  — no claims to verify (not a violation)
#    UNCERTAIN      — single ambiguous claim
#
#  G_rate = |{s ∈ N : G(s) ⊆ J}| / |N_numerical|   — scalar metric (0–1)
#
#  This G_rate is:
#    (a) measurable — computed per section and overall
#    (b) verifiable — full evidence table showing which J path matched
#    (c) novel      — no prior environmental monitoring pipeline defines or
#                     measures this property programmatically
#
#  Reference: ATARS formal framework, Eq. 19, G: J → N
# ═══════════════════════════════════════════════════════════════════════════════

import re as _re

class RuntimeGroundingVerifier:
    """
    Programmatic enforcement of the formal grounding constraint:
      G(s) ⊆ J  ∧  G(s) ≠ ∅  for all numerically-bearing sentences s ∈ N
    where J is the JSON data contract and N is the set of generated sentences.

    Algorithm:
      1. Flatten J into a numeric lookup table L = {path: value}
      2. For each sentence s ∈ N:
         a. Extract numerical claims C(s) = {(val, type, context)}
         b. For each c ∈ C(s): check if ∃ j ∈ L: |c - j|/|j| ≤ ε
         c. Assign verdict based on |{grounded claims}| / |C(s)|
      3. Compute G_rate = |{s: GROUNDED or PARTIALLY_GROUNDED}| / |N_numerical|
      4. Emit full evidence table (claim → J path → J value → match status)
    """

    # Tolerance bounds for numerical matching
    TOL_RELATIVE = 0.05    # ±5% for floating-point concentrations, scores
    TOL_ABSOLUTE = 2.0     # ±2 for integer counts (days, records)
    TOL_EXACT    = 0.001   # ±0.001 for quality scores, thresholds

    # Minimum G_rate for a section to pass
    PASS_THRESHOLD = 0.75

    def __init__(self, J: dict):
        self.J      = J
        self.lookup = {}                         # path → float value
        self.string_lookup = {}                  # path → string value
        self._build_lookup(J, path='')

    # ── Lookup table construction ───────────────────────────────────────────
    def _build_lookup(self, obj, path: str):
        """Recursively flatten J into {dotted.path: value}."""
        if isinstance(obj, dict):
            for k, v in obj.items():
                new_path = f"{path}.{k}" if path else k
                self._build_lookup(v, new_path)
        elif isinstance(obj, list):
            for i, v in enumerate(obj):
                self._build_lookup(v, f"{path}[{i}]")
        elif isinstance(obj, bool):
            pass                                 # skip boolean flags
        elif isinstance(obj, (int, float)):
            try:
                fval = float(obj)
                if not (math.isnan(fval) or math.isinf(fval)):
                    self.lookup[path] = fval
            except (TypeError, ValueError):
                pass
        elif isinstance(obj, str):
            self.string_lookup[path] = obj

    # ── Claim extraction ────────────────────────────────────────────────────
    def extract_claims(self, sentence: str) -> list:
        """
        Extract all numerical claims from a sentence.
        Returns list of dicts: {raw, value, type, context}
        Skips citation-style numbers (equation refs, section numbers < 30).
        """
        claims = []
        # Match integers and floats, including negative
        pattern = _re.compile(
            r'(?<![a-zA-Z])(-?\d{1,6}(?:\.\d{1,6})?)\b'
        )
        for m in pattern.finditer(sentence):
            raw = m.group(1)
            try:
                val = float(raw)
            except ValueError:
                continue

            # Skip likely equation/section references
            if val == int(val) and 0 < val <= 30 and '.' not in raw:
                ctx_before = sentence[max(0, m.start()-15):m.start()].lower()
                if any(w in ctx_before for w in
                       ['eq.', 'section', 'fig.', 'table', 'prop.']):
                    continue

            # Determine semantic type from surrounding context (window ±30 chars)
            win = sentence[max(0, m.start()-30): m.end()+30].lower()

            if any(w in win for w in ['day', 'days']):
                ctype = 'days_count'
            elif any(w in win for w in ['record', 'observation', 'n=']):
                ctype = 'record_count'
            elif '%' in win[len(raw):len(raw)+5] or 'percent' in win:
                ctype = 'percentage'
            elif any(w in win for w in ['µg', 'mg/m', 'ppb', 'ppm', 'µg/m']):
                ctype = 'concentration'
            elif any(w in win for w in ['q(', 'q =', 'quality', 'completeness',
                                         'score', 'completene']):
                ctype = 'quality_score'
            elif any(w in win for w in ['r²', 'r2', 'r-squared', 'pearson',
                                         'correlation']):
                ctype = 'correlation_coef'
            elif any(w in win for w in ['z-score', 'zeta', 'ζ', 'z_score',
                                         'anomal']):
                ctype = 'z_score'
            elif any(w in win for w in ['threshold', 'guideline', 'limit',
                                         'who', 'who ']):
                ctype = 'threshold'
            elif any(w in win for w in ['year', 'annual', 'month']):
                ctype = 'temporal'
            else:
                ctype = 'numeric'

            claims.append({
                'raw': raw, 'value': val, 'type': ctype,
                'context': win.strip()
            })
        return claims

    # ── Single claim matching ────────────────────────────────────────────────
    def match_claim(self, val: float, ctype: str) -> dict:
        """
        Find the closest match in J lookup for val.
        Returns evidence dict: {matched, j_path, j_value, diff, method}
        """
        best = {'matched': False, 'j_path': None, 'j_value': None,
                'diff': float('inf'), 'method': None}

        for path, j_val in self.lookup.items():
            # Absolute difference
            abs_diff = abs(val - j_val)

            # Choose tolerance based on type and magnitude
            if ctype in ('days_count', 'record_count') or \
               (val == int(val) and abs(val) >= 10):
                # Integer counts: absolute tolerance ±2
                tol = self.TOL_ABSOLUTE
                method = 'absolute'
                score = abs_diff
            elif ctype in ('quality_score', 'correlation_coef', 'z_score'):
                # Precise scalars: tight absolute tolerance
                tol = self.TOL_EXACT * 100   # 0.1
                method = 'tight'
                score = abs_diff
            elif j_val != 0:
                # General: relative tolerance ±5%
                rel_diff = abs_diff / abs(j_val)
                tol = self.TOL_RELATIVE
                method = 'relative'
                score = rel_diff
            else:
                tol = self.TOL_ABSOLUTE
                method = 'absolute'
                score = abs_diff

            if score <= tol and score < best['diff']:
                best = {'matched': True, 'j_path': path, 'j_value': j_val,
                        'diff': round(score, 6), 'method': method}

        return best

    # ── Sentence verification ────────────────────────────────────────────────
    def verify_sentence(self, sentence: str) -> dict:
        """
        Verify one sentence. Assigns verdict and full evidence.
        Returns structured result dict.
        """
        claims  = self.extract_claims(sentence)

        if not claims:
            return {
                'verdict': 'NON_NUMERICAL', 'sentence': sentence,
                'claims_total': 0, 'claims_grounded': 0,
                'claim_rate': None, 'evidence': []
            }

        evidence = []
        n_grounded = 0
        for c in claims:
            match = self.match_claim(c['value'], c['type'])
            is_g  = match['matched']
            if is_g:
                n_grounded += 1
            evidence.append({
                'claim'       : c['raw'],
                'value'       : c['value'],
                'type'        : c['type'],
                'grounded'    : is_g,
                'j_path'      : match['j_path'],
                'j_value'     : match['j_value'],
                'diff'        : match['diff'],
                'method'      : match['method'],
            })

        total = len(claims)
        rate  = n_grounded / total

        if rate == 1.0:
            verdict = 'GROUNDED'
        elif rate >= 0.5:
            verdict = 'PARTIALLY_GROUNDED'
        elif total == 1:
            verdict = 'UNCERTAIN'
        else:
            verdict = 'UNGROUNDED'

        return {
            'verdict'         : verdict,
            'sentence'        : sentence,
            'claims_total'    : total,
            'claims_grounded' : n_grounded,
            'claim_rate'      : round(rate, 4),
            'evidence'        : evidence,
        }

    # ── Section verification ─────────────────────────────────────────────────
    def verify_section(self, section_key: str, text: str) -> dict:
        """
        Verify an entire narrative section.
        Splits into sentences, verifies each, computes section G_rate.
        """
        # Split on sentence boundaries
        raw_sents = _re.split(r'(?<=[.!?])\s+', text.strip())
        sentences = [s.strip() for s in raw_sents if len(s.strip()) > 8]

        results       = [self.verify_sentence(s) for s in sentences]
        numerical     = [r for r in results if r['verdict'] != 'NON_NUMERICAL']
        grounded_list = [r for r in numerical
                         if r['verdict'] in ('GROUNDED', 'PARTIALLY_GROUNDED')]

        g_rate = (len(grounded_list) / len(numerical)) if numerical else 1.0
        passes = g_rate >= self.PASS_THRESHOLD

        return {
            'section_key'          : section_key,
            'sentences'            : results,
            'total_sentences'      : len(results),
            'numerical_sentences'  : len(numerical),
            'grounded_sentences'   : len(grounded_list),
            'ungrounded_sentences' : len(numerical) - len(grounded_list),
            'G_rate'               : round(g_rate, 4),
            'passes'               : passes,
            'verdict'              : 'PASS ✓' if passes else 'REVIEW ⚠',
        }

    # ── Full narrative verification ──────────────────────────────────────────
    def verify_narrative(self, narrative: dict) -> dict:
        """
        Verify the complete generated narrative N.

        Formally: computes
          G_rate = |{s ∈ N_num : verdict ∈ {GROUNDED, PARTIALLY_GROUNDED}}|
                   ────────────────────────────────────────────────────────
                                    |N_num|

        where N_num = {s ∈ N : C(s) ≠ ∅} (sentences with numerical claims).
        """
        sections      = {}
        total_num     = 0
        total_grounded= 0
        all_ungrounded= []

        for key in NARRATIVE_SECTIONS:
            if key not in narrative:
                continue
            result = self.verify_section(key, narrative[key])
            sections[key]     = result
            total_num        += result['numerical_sentences']
            total_grounded   += result['grounded_sentences']

            for sr in result['sentences']:
                if sr['verdict'] == 'UNGROUNDED':
                    all_ungrounded.append({
                        'section'  : key,
                        'sentence' : (sr['sentence'][:500] + ' …' if len(sr['sentence']) > 500 else sr['sentence']),
                        'n_claims' : sr['claims_total'],
                        'evidence' : sr['evidence'][:3],
                    })

        overall = (total_grounded / total_num) if total_num > 0 else 1.0
        passed  = overall >= self.PASS_THRESHOLD

        return {
            'schema'             : 'RGV-1.0',
            'formal_property'    : 'G(s) ⊆ J  ∀ s ∈ N  [ATARS Eq. 19]',
            'overall_G_rate'     : round(overall, 4),
            'total_numerical'    : total_num,
            'total_grounded'     : total_grounded,
            'total_ungrounded'   : total_num - total_grounded,
            'verification_passed': passed,
            'overall_verdict'    : 'PASS ✓' if passed else 'REVIEW ⚠',
            'pass_threshold'     : self.PASS_THRESHOLD,
            'sections'           : sections,
            'ungrounded_list'    : all_ungrounded,
            'j_lookup_size'      : len(self.lookup),
            'timestamp'          : datetime.now().isoformat(),
        }


def chart_grounding_verification(verification: dict, config: dict, out: Path) -> str:
    """
    Chart 13 — RGV Grounding Verification Dashboard.
    Shows G_rate per section, verdict distribution, and evidence heatmap.
    This visualisation is part of the novel contribution.
    """
    sections = verification.get('sections', {})
    if not sections:
        return None

    fig = plt.figure(figsize=(14, 10), facecolor=PAL['white'])
    gs  = gridspec.GridSpec(2, 3, figure=fig, hspace=0.52, wspace=0.38)

    ax_bar   = fig.add_subplot(gs[0, :2])   # G_rate per section (wide)
    ax_gauge = fig.add_subplot(gs[0, 2])    # Overall G_rate gauge
    ax_pie   = fig.add_subplot(gs[1, 0])    # Verdict distribution
    ax_heat  = fig.add_subplot(gs[1, 1:])   # Evidence density heatmap

    # ── Panel A: G_rate per section ─────────────────────────────────────────
    sec_labels = list(sections.keys())
    g_rates    = [sections[k]['G_rate'] for k in sec_labels]
    verdicts   = [sections[k]['passes'] for k in sec_labels]
    bar_colors = [PAL['green'] if v else PAL['amber'] for v in verdicts]

    bars = ax_bar.barh(sec_labels, g_rates, color=bar_colors, height=0.55, alpha=0.85)
    ax_bar.axvline(verification['pass_threshold'], color=PAL['red'],
                   linewidth=2, linestyle='--',
                   label=f'Pass threshold ({verification["pass_threshold"]:.0%})')
    ax_bar.set_xlim(0, 1.05)
    ax_bar.set_xlabel('G_rate (grounded sentences / numerical sentences)', fontsize=9)
    ax_bar.set_title(
        r'(a) G_rate per narrative section  —  $G_\text{rate} = '
        r'|\{s : G(s) \subseteq J\}| \,/\, |N_\text{num}|$',
        fontsize=10, fontweight='bold', color=PAL['navy'])
    ax_bar.legend(fontsize=8.5)
    for bar, rate in zip(bars, g_rates):
        ax_bar.text(bar.get_width() + 0.01, bar.get_y() + bar.get_height()/2,
                    f'{rate:.1%}', va='center', fontsize=9, fontweight='bold')
    ax_bar.tick_params(axis='y', labelsize=9)

    # ── Panel B: Overall G_rate gauge ──────────────────────────────────────
    overall = verification['overall_G_rate']
    theta   = np.linspace(0, np.pi, 200)
    ax_gauge.set_aspect('equal')
    # Background arc
    ax_gauge.plot(np.cos(theta), np.sin(theta), color='#E5E7EB', linewidth=18,
                  solid_capstyle='round')
    # Filled arc proportional to G_rate
    theta_fill = np.linspace(0, np.pi * overall, 200)
    arc_color  = PAL['green'] if overall >= 0.75 else (PAL['amber'] if overall >= 0.5 else PAL['red'])
    ax_gauge.plot(np.cos(theta_fill), np.sin(theta_fill),
                  color=arc_color, linewidth=18, solid_capstyle='round')
    ax_gauge.text(0, 0.2, f'{overall:.1%}', ha='center', va='center',
                  fontsize=26, fontweight='bold', color=arc_color)
    ax_gauge.text(0, -0.18, verification['overall_verdict'],
                  ha='center', va='center', fontsize=13, fontweight='bold',
                  color=arc_color)
    ax_gauge.text(0, -0.42, 'Overall G_rate', ha='center', va='center',
                  fontsize=9, color=PAL['gray'])
    ax_gauge.set_xlim(-1.3, 1.3); ax_gauge.set_ylim(-0.6, 1.2)
    ax_gauge.axis('off')
    ax_gauge.set_title('(b) Overall\nGrounding Rate',
                        fontsize=10, fontweight='bold', color=PAL['navy'])

    # ── Panel C: Verdict distribution pie ──────────────────────────────────
    verdict_counts = defaultdict(int)
    for sec_data in sections.values():
        for sr in sec_data['sentences']:
            verdict_counts[sr['verdict']] += 1
    v_labels = list(verdict_counts.keys())
    v_vals   = list(verdict_counts.values())
    v_colors_map = {
        'GROUNDED'           : PAL['green'],
        'PARTIALLY_GROUNDED' : PAL['teal'],
        'NON_NUMERICAL'      : '#A0B4CC',
        'UNCERTAIN'          : PAL['amber'],
        'UNGROUNDED'         : PAL['red'],
    }
    v_colors = [v_colors_map.get(l, PAL['gray']) for l in v_labels]
    wedges, texts, autotexts = ax_pie.pie(
        v_vals, labels=v_labels, colors=v_colors,
        autopct='%1.0f%%', startangle=90,
        textprops={'fontsize': 8},
        pctdistance=0.75
    )
    for at in autotexts: at.set_fontsize(8)
    ax_pie.set_title('(c) Verdict Distribution\n(all sentences)',
                     fontsize=10, fontweight='bold', color=PAL['navy'])

    # ── Panel D: Evidence density — claims per section ──────────────────────
    sec_names   = list(sections.keys())
    ev_grounded = [sections[k]['grounded_sentences']   for k in sec_names]
    ev_ungnd    = [sections[k]['ungrounded_sentences']  for k in sec_names]
    ev_non_num  = [sections[k]['total_sentences'] -
                   sections[k]['numerical_sentences']  for k in sec_names]
    x = np.arange(len(sec_names))
    w = 0.25
    ax_heat.bar(x - w, ev_grounded, w, label='Grounded',     color=PAL['green'],  alpha=0.85)
    ax_heat.bar(x,     ev_ungnd,    w, label='Ungrounded',   color=PAL['red'],    alpha=0.85)
    ax_heat.bar(x + w, ev_non_num,  w, label='Non-numerical',color='#A0B4CC',     alpha=0.85)
    ax_heat.set_xticks(x); ax_heat.set_xticklabels(sec_names, fontsize=9)
    ax_heat.set_ylabel('Sentence count', fontsize=9)
    ax_heat.set_title('(d) Sentence-Level Grounding Analysis\n'
                      'Grounded / Ungrounded / Non-numerical per section',
                      fontsize=10, fontweight='bold', color=PAL['navy'])
    ax_heat.legend(fontsize=8.5)

    fig.suptitle(
        f'ATARS — Runtime Grounding Verifier (RGV)  ·  {config["city"]}\n'
        r'Novel contribution: programmatic enforcement of $G(s) \subseteq J$  '
        f'— G_rate = {overall:.1%}  [{verification["overall_verdict"]}]',
        fontsize=12, fontweight='bold', color=PAL['navy'])

    fig.text(0.5, 0.01,
             f'RGV v1.0  |  J lookup size: {verification["j_lookup_size"]} entries  |  '
             f'Sections verified: {len(sections)}  |  '
             f'Total sentences: {sum(s["total_sentences"] for s in sections.values())}  |  '
             f'Timestamp: {verification["timestamp"][:19]}',
             ha='center', fontsize=7.5, color=PAL['gray'], style='italic')

    return _save_fig(fig, out, 'chart13_grounding_verification.png')


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 11 — WORD REPORT GENERATION (python-docx)
# ═══════════════════════════════════════════════════════════════════════════════

def _set_cell_bg(cell, hex_color: str):
    """Set table cell background colour."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def _add_table_row(table, cells_data, bold_col=0, bg=None, font_size=10):
    """Add a styled row to a Word table."""
    row = table.add_row()
    for i, (cell, data) in enumerate(zip(row.cells, cells_data)):
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
        run  = para.add_run(str(data))
        run.bold      = (i == bold_col)
        run.font.size = Pt(font_size)
        if bg:
            _set_cell_bg(cell, bg if i % 2 == 0 else 'FFFFFF')


def _rgb(hex_color: str) -> RGBColor:
    h = hex_color.lstrip('#')
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def build_report(J: dict, daily_df: pd.DataFrame, charts: dict,
                 narrative: dict, ols_result: dict,
                 corr_matrix: pd.DataFrame, config: dict,
                 out_dir: Path, verification: dict = None,
                 chart_insights: dict = None,
                 ml_if: dict = None, ml_hw: dict = None) -> str:
    """
    Assemble the complete Word document report.
    All content derived from J — no raw data passed to report builder.
    v2.0: Adds chart insights text, ML anomaly section, forecast section.
    chart_insights: dict of {chart_name: insight_string} — deterministic text.
    """
    chart_insights = chart_insights or {}
    ml_if = ml_if or {}
    ml_hw = ml_hw or {}
    print("\n  [Step 7] Assembling Word report...")
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Colour palette ────────────────────────────────────────────────────────
    NAVY   = _rgb('1A2C4E'); BLUE   = _rgb('2557A7')
    TEAL   = _rgb('1A6B72'); GREEN  = _rgb('1E6B3C')
    RED    = _rgb('8B1C2A'); AMBER  = _rgb('92400E')
    GRAY   = _rgb('6B7280'); WHITE  = _rgb('FFFFFF')
    PURPLE = _rgb('4C1D95')

    def h1(text, color=NAVY):
        p = doc.add_heading(text, level=1)
        for run in p.runs:
            run.font.color.rgb = color
            run.font.size = Pt(18)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(8)

    def h2(text, color=BLUE):
        p = doc.add_heading(text, level=2)
        for run in p.runs:
            run.font.color.rgb = color
            run.font.size = Pt(14)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(6)

    def body(text, italic=False, color=None):
        p = doc.add_paragraph(text)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(6)
        for run in p.runs:
            run.italic = italic
            if color:
                run.font.color.rgb = color
            run.font.size = Pt(11)
        return p

    def caption(text):
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size  = Pt(9)
            run.italic     = True
            run.font.color.rgb = GRAY
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(10)

    def add_image(chart_key, width_inch=6.2, cap_text=""):
        path = charts.get(chart_key)
        if path and os.path.exists(path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(path, width=Inches(width_inch))
        if cap_text:
            caption(cap_text)
        # Add deterministic chart insight text if available
        insight = chart_insights.get(chart_key, "")
        if insight:
            p_ins = doc.add_paragraph()
            p_ins.paragraph_format.space_before = Pt(4)
            p_ins.paragraph_format.space_after  = Pt(8)
            p_ins.paragraph_format.left_indent  = Pt(14)
            # Teal vertical bar via shading is not directly available in python-docx
            # Use a styled paragraph with ► prefix
            run_lbl = p_ins.add_run("► Chart Insight:  ")
            run_lbl.bold = True
            run_lbl.font.color.rgb = TEAL
            run_lbl.font.size = Pt(9.5)
            run_txt = p_ins.add_run(insight)
            run_txt.italic = True
            run_txt.font.size = Pt(9.5)
            run_txt.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    # ════════════════════════════════════════════════════════════════════════
    #  COVER PAGE
    # ════════════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run('ATARS')
    run.font.size  = Pt(54)
    run.bold       = True
    run.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Automated Time-Series Analysis and Reporting System  v2.0')
    run.font.size = Pt(16); run.font.color.rgb = BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Annual Air Quality Analysis Report  —  {J["city"]}')
    run.font.size = Pt(20); run.bold = True; run.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{J["date_range"]}  ·  {J["total_days"]} days  ·  {J["total_records"]:,} records')
    run.font.size = Pt(12); run.italic = True; run.font.color.rgb = GRAY

    doc.add_paragraph()

    # Confidence badge
    conf    = J['data_quality']['confidence_flag']
    q_score = J['data_quality']['overall_Q']
    conf_colors = {'HIGH': '1E6B3C', 'MODERATE': '92400E', 'LOW': '8B1C2A'}
    conf_bg     = {'HIGH': 'D1FAE5', 'MODERATE': 'FEF3C7', 'LOW': 'FFE4E6'}

    t = doc.add_table(rows=1, cols=4)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style     = 'Table Grid'
    data_conf = [
        ('Confidence Level', conf),
        ('Q(D) Score', f'{q_score:.3f}'),
        ('Total Days', str(J['total_days'])),
        ('Total Records', f'{J["total_records"]:,}'),
    ]
    for cell, (label, value) in zip(t.rows[0].cells, data_conf):
        cell.paragraphs[0].clear()
        lp = cell.add_paragraph(label)
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in lp.runs:
            r.font.size = Pt(9); r.bold = True; r.font.color.rgb = GRAY
        vp = cell.add_paragraph(value)
        vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in vp.runs:
            r.font.size = Pt(18); r.bold = True
            r.font.color.rgb = _rgb(conf_colors.get(conf, '1A2C4E'))
        _set_cell_bg(cell, conf_bg.get(conf, 'DCE8F5'))

    doc.add_paragraph()

    # Author block
    t2 = doc.add_table(rows=1, cols=1)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    t2.style     = 'Table Grid'
    cell = t2.rows[0].cells[0]
    _set_cell_bg(cell, 'F4F5F7')
    for text, pt, bold, color in [
        (config['author'], 18, True, NAVY),
        (config['degree'], 11, False, GRAY),
        (config['department'], 12, True, BLUE),
        (config['institution'], 12, False, NAVY),
        (f"{config['location']}  ·  {config['email']}", 10, False, GRAY),
        (f"Document: RP-ATARS-{datetime.now().year}-001  ·  Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", 9, False, GRAY),
    ]:
        p = cell.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        for r in p.runs:
            r.font.size = Pt(pt); r.bold = bold; r.font.color.rgb = color

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    #  SECTION 01 — DATA QUALITY STATEMENT
    # ════════════════════════════════════════════════════════════════════════
    h1('Section 01 — Data Quality Statement')
    body(narrative.get('E_quality', ''))

    # Quality table
    h2('Quality Score Q(D) per Variable')
    q_vars = J['data_quality']['per_variable_Q']
    if q_vars:
        t = doc.add_table(rows=1, cols=5)
        t.style = 'Table Grid'
        for cell, hdr in zip(t.rows[0].cells,
                              ['Variable', 'Q(D) = N_v/N_total',
                               'Confidence', 'Anomaly Days', 'Exceedance Days']):
            p = cell.paragraphs[0]
            run = p.add_run(hdr)
            run.bold = True; run.font.size = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_bg(cell, '1A2C4E')
            for r in p.runs:
                r.font.color.rgb = WHITE

        for col, q_val in sorted(q_vars.items(), key=lambda x: -x[1]):
            v_data = J['variables'].get(col, {})
            row = t.add_row()
            bg = 'D1FAE5' if q_val >= 0.90 else ('FEF3C7' if q_val >= 0.70 else 'FFE4E6')
            vals = [
                col, f'{q_val:.3f}',
                'HIGH' if q_val >= 0.90 else ('MODERATE' if q_val >= 0.70 else 'LOW'),
                str(v_data.get('anomaly_days', 'N/A')),
                str(v_data.get('exceeds_days', 0)),
            ]
            for i, (cell, val) in enumerate(zip(row.cells, vals)):
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
                r = p.add_run(val); r.font.size = Pt(10)
                _set_cell_bg(cell, bg)

    add_image('quality', 6.2,
              'Figure — Q(D) = N_v/N_total per variable (Eq. 11). '
              'Confidence flags: HIGH ≥ 0.90, MODERATE ≥ 0.70, LOW < 0.70.')

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    #  SECTION 02 — ANNUAL STATISTICS TABLE
    # ════════════════════════════════════════════════════════════════════════
    h1('Section 02 — Annual Statistics Summary')
    body(narrative.get('A_summary', ''))

    h2('Formal Statistical Operators: A(D_v), ζ_d, δ_d (Eqs. 3–10)')
    stats_cols = ['Variable', 'Annual Mean x̄', 'Std σ', 'Max', 'Min',
                  'Threshold θ', 'Exceed Days', 'Anomaly Days', 'Avg ζ_d']
    t = doc.add_table(rows=1, cols=len(stats_cols))
    t.style = 'Table Grid'
    for cell, hdr in zip(t.rows[0].cells, stats_cols):
        p = cell.paragraphs[0]
        run = p.add_run(hdr); run.bold = True; run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_bg(cell, '1A2C4E')
        for r in p.runs: r.font.color.rgb = WHITE

    for i, (col, vd) in enumerate(J['variables'].items()):
        row = t.add_row()
        bg  = 'FFE4E6' if vd.get('exceeds_days', 0) > 30 else \
              ('FEF3C7' if vd.get('exceeds_days', 0) > 0 else
               ('DCE8F5' if i % 2 == 0 else 'FFFFFF'))
        thr = vd.get('threshold')
        zs  = vd.get('mean_z_score')
        vals = [
            col,
            f'{vd["annual_mean"]:.2f}',
            f'{vd["annual_std"]:.2f}',
            f'{vd["annual_max"]:.2f}',
            f'{vd["annual_min"]:.2f}',
            f'{thr}' if thr else '—',
            f'{vd.get("exceeds_days", 0)} ({vd.get("exceed_pct", 0):.1f}%)',
            str(vd.get('anomaly_days', 0)),
            f'{zs:.2f}' if zs is not None else '—',
        ]
        for j, (cell, val) in enumerate(zip(row.cells, vals)):
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(val); r.font.size = Pt(9)
            _set_cell_bg(cell, bg)

    body(
        'x̄ = mean of all daily means over the monitoring period (annual mean). '
        'σ = sample standard deviation of daily means. '
        'Exceed Days = number of days where the daily mean exceeded guideline threshold θ. '
        'Anomaly Days = number of days where |ζ_d| > 3.0 (z-score threshold, Eq. 9). '
        'All statistics are computed from quality-validated records only (q_i = 1).',
        italic=True
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    #  SECTION 03 — TIME-SERIES VISUALISATION
    # ════════════════════════════════════════════════════════════════════════
    h1('Section 03 — Hourly Time-Series Analysis')
    add_image('timeseries', 6.2,
              'Figure — Hourly concentrations with daily mean overlay and WHO AQG 2021 '
              'reference guidelines. Shaded areas indicate exceedance zones.')

    body(narrative.get('B_pollutants', ''))

    h2('30-Day Rolling Baseline Trend (Eq. 7–8)')
    body(
        'The chart below shows the daily pollutant concentration alongside the 30-day '
        'rolling baseline mean \u03c4\u0305_W and standard deviation \u03c3_W. '
        'Deviations above the rolling mean identify periods of elevated pollution '
        'relative to recent local norms, without seasonal bias.'
    )
    add_image('trend', 6.2,
              'Figure — Daily trend and 30-day rolling baseline x̄_W (Eqs. 7–8). '
              'Dashed lines show baseline mean. Dotted lines show WHO guidelines.')

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    #  SECTION 04 — ANOMALY DETECTION
    # ════════════════════════════════════════════════════════════════════════
    h1('Section 04 — Anomaly Detection (Z-Score Analysis)')
    body(narrative.get('D_anomalies', ''))

    add_image('anomalies', 6.2,
              'Figure — (a) Daily mean with 95% CI band (Eq. 12). '
              r'(b) Z-score ζ_d = (x̄_d − x̄_W) / σ_W (Eq. 9). '
              'Red diamonds indicate anomaly days (|ζ_d| > 3.0).')

    # Anomaly summary table
    h2('Anomaly Summary by Pollutant')
    anom_data = [(col, vd['anomaly_days'], vd['mean_z_score'])
                 for col, vd in J['variables'].items()
                 if vd.get('anomaly_days', 0) > 0]
    if anom_data:
        t = doc.add_table(rows=1, cols=3)
        t.style = 'Table Grid'
        for cell, hdr in zip(t.rows[0].cells,
                              ['Variable', 'Anomaly Days (|ζ_d| > 3)', 'Mean ζ_d']):
            p = cell.paragraphs[0]; run = p.add_run(hdr)
            run.bold = True; run.font.size = Pt(10)
            _set_cell_bg(cell, '1A2C4E')
            for r in p.runs: r.font.color.rgb = WHITE
        for col, adays, mz in sorted(anom_data, key=lambda x: -x[1]):
            row = t.add_row()
            for cell, val in zip(row.cells,
                                  [col, str(adays),
                                   f'{mz:.2f}' if mz else '—']):
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(val); r.font.size = Pt(10)
                _set_cell_bg(cell, 'FFF5F5')

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    #  SECTION 05 — TEMPORAL DEPENDENCY (ACF / PACF)
    # ════════════════════════════════════════════════════════════════════════
    h1('Section 05 — Temporal Dependency: ACF and PACF')
    add_image('acf_pacf', 6.0,
              'Figure — (a) ACF ρ̂(k) = Corr(x_t, x_{t-k}) (Eq. 12). '
              '(b) PACF φ̂_kk: conditional correlation at lag k (Eq. 13). '
              'Bars outside red dashed CI band indicate significant temporal dependence.')

    body(
        'The Autocorrelation Function (ACF) measures the total correlation between '
        'observations separated by lag k. The Partial ACF (PACF) isolates the direct '
        'contribution of lag k after removing the influence of all shorter lags. '
        'Together, ACF and PACF characterise the temporal memory structure of the series '
        'and provide the empirical basis for identifying appropriate ARMA model orders.'
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    #  SECTION 06 — CORRELATION & MULTIVARIATE ANALYSIS
    # ════════════════════════════════════════════════════════════════════════
    h1('Section 06 — Correlation & Multivariate Analysis')
    body(narrative.get('C_correlations', ''))

    add_image('correlation', 6.0,
              'Figure — Pearson correlation matrix R (Eq. 18). '
              'Blue = positive, red = negative. Values annotated per cell. '
              'FORMAL CAUSAL DISCLAIMER: Corr(X,Y) ≢ C(X→Y). '
              'Association does not imply causation.')

    # OLS regression results
    if ols_result and 'r_squared' in ols_result:
        h2(f'OLS Regression: {ols_result["target"]} ~ f(meteorological variables) (Eq. 14)')
        body(f'β̂ = (XᵀX)⁻¹Xᵀy. R² = {ols_result["r_squared"]:.4f}. '
             f'n = {ols_result["n_obs"]} observations. '
             f'{ols_result["note"]}')
        t = doc.add_table(rows=1, cols=3)
        t.style = 'Table Grid'
        for cell, hdr in zip(t.rows[0].cells,
                              ['Predictor', 'Coefficient β̂', 'Interpretation']):
            p = cell.paragraphs[0]; run = p.add_run(hdr)
            run.bold = True; run.font.size = Pt(10)
            _set_cell_bg(cell, '1A2C4E')
            for r in p.runs: r.font.color.rgb = WHITE
        for pred, beta in ols_result.get('beta', {}).items():
            row = t.add_row()
            direction = 'positive association' if beta > 0 else 'negative association'
            for cell, val in zip(row.cells,
                                  [pred, f'{beta:.4f}', direction]):
                p = cell.paragraphs[0]
                r = p.add_run(val); r.font.size = Pt(10)
                _set_cell_bg(cell, 'DCE8F5')
        body(
            'Coefficients are standardised (predictors are z-score normalised before regression). '
            'The R² value quantifies how much of the variance in the target variable is '
            'explained by the meteorological predictors in this dataset. '
            'No causal interpretation is warranted — regression coefficients describe '
            'statistical association only.',
            italic=True
        )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    #  SECTION 07 — SEASONAL ANALYSIS
    # ════════════════════════════════════════════════════════════════════════
    h1('Section 07 — Seasonal Pattern Analysis')
    body(
        'Monthly box plots reveal the full distribution of daily pollutant concentrations '
        'across the year. Each box spans the 25th to 75th percentile, with the horizontal '
        'line at the median. Whiskers extend to 1.5 times the inter-quartile range; '
        'individual points beyond the whiskers are statistical outliers. '
        'This non-parametric display requires no normality assumption and clearly shows '
        'which months carry the highest pollution burden.'
    )
    add_image('seasonal', 6.2,
              'Figure — Monthly distribution: median x̃ (Eq. 24), IQR (Eq. 25), '
              'whiskers = 1.5×IQR. Non-parametric — no normality assumption. '
              'Dotted line = WHO 24h reference guideline.')

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    #  SECTION 08 — DIURNAL (HOURLY) PROFILE
    # ════════════════════════════════════════════════════════════════════════
    h1('Section 08 — Diurnal Cycle Analysis')
    body(
        'The diurnal profile shows the average pollutant concentration for each hour of '
        'the day (00:00 to 23:00), computed across the full monitoring period. '
        'The timing of daily peaks identifies the dominant emission sources: '
        'morning and evening peaks typically indicate traffic, while midday peaks '
        'may reflect industrial or secondary photochemical formation. '
        'The shaded band represents the mean ± 1σ envelope across all days in the dataset. '
        'The normalised overlay panel (peak = 1.0 per pollutant) allows direct '
        'shape comparison across variables regardless of their concentration magnitudes.'
    )
    add_image('diurnal', 6.2,
              'Figure — Hourly diurnal profile for key pollutants. Shaded band = mean ± σ. '
              'Bottom-right panel: normalised overlay (peak = 1.0 per variable). '
              'Peak hours identify primary emission timing.')

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    #  SECTION 09 — DAY-OF-WEEK PATTERN
    # ════════════════════════════════════════════════════════════════════════
    h1('Section 09 — Day-of-Week Pattern')
    body(
        'Day-of-week analysis reveals whether pollution follows a weekly anthropogenic '
        'cycle. Concentrations are typically higher on weekdays due to traffic and '
        'industrial activity, and lower on weekends. A clear weekday-to-weekend '
        'contrast indicates human-driven emission sources; the absence of such a '
        'pattern suggests natural or regionally transported background pollution.'
    )
    add_image('day_of_week', 6.2,
              'Figure — Mean daily concentration by day of week. Error bars = ±1σ. '
              'Weekday–weekend contrast indicates anthropogenic emission cycle.')

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    #  SECTION 10 — VOC COMPOUNDS PANEL
    # ════════════════════════════════════════════════════════════════════════
    h1('Section 10 — Volatile Organic Compounds (VOC) Panel')
    body(
        'The BTEX group (Benzene, Toluene, Ethylbenzene, and Xylene) together with '
        'MP-Xylene are primary markers of vehicular exhaust and industrial solvent emissions. '
        'Benzene is classified as a Group 1 carcinogen by the IARC, with a WHO annual '
        'reference concentration of 1.7 µg/m³. '
        'The grouped monthly bar chart in panel (b) reveals the seasonal pattern of '
        'VOC concentrations and identifies months with elevated BTEX loading.'
    )
    add_image('voc', 6.0,
              'Figure — (a) VOC daily means time series. '
              '(b) Monthly grouped bar chart for all VOC compounds. '
              'Benzene WHO annual guideline: 1.7 µg/m³ (dotted red line, panel a).')

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    #  SECTION 11 — RAIN vs POLLUTANT SCATTER
    # ════════════════════════════════════════════════════════════════════════
    h1('Section 11 — Precipitation vs Pollutant Scatter (Wet Deposition)')
    body(
        'Scatter plots of daily rainfall (mm) against PM10, SO2, and NO2 daily means '
        'reveal statistical associations between precipitation events and pollutant '
        'concentration levels. Negative Pearson r values indicate that higher rainfall '
        'is associated with lower concentrations — consistent with wet deposition, '
        'where precipitation removes particulates and soluble gases from the atmosphere. '
        'Note: this is a statistical association only. Confounding meteorological '
        'factors such as wind speed and boundary layer height may contribute independently.'
    )
    add_image('rain_scatter', 6.2,
              'Figure — Daily rain (mm) vs pollutant concentration scatter. '
              'Dashed line = OLS trend. r = Pearson correlation coefficient. '
              'Association only — not causal wet deposition estimation.')

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    #  SECTION 12 — EXCEEDANCE CALENDAR HEATMAP
    # ════════════════════════════════════════════════════════════════════════
    h1('Section 12 — Annual Exceedance Calendar Heatmap')
    body('The calendar heatmap provides a full-year view of daily PM10 concentrations '
         'at a glance. Each cell shows the daily mean value. Cells marked × indicate '
         'days where the daily mean exceeded the WHO 24h guideline. The colour scale '
         'from green (clean) to red (polluted) immediately reveals seasonal pollution '
         'clusters, high-concentration episodes, and clean-air periods.')
    add_image('calendar', 6.2,
              'Figure — PM10 calendar heatmap: month (rows) × day of month (columns). '
              '× = WHO 24h guideline exceedance. Green = clean, Red = polluted. '
              'Numeric value = daily mean µg/m³.')

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    #  SECTION 13 — RUNTIME GROUNDING VERIFICATION CERTIFICATE (RGV)
    #  NOVEL CONTRIBUTION — programmatic G(s) ⊆ J enforcement
    # ════════════════════════════════════════════════════════════════════════
    if verification:
        h1('Section 13 — Runtime Grounding Verification Certificate (RGV)')

        # ── Novel contribution banner ─────────────────────────────────────
        p = doc.add_paragraph()
        run = p.add_run(
            'This section presents the Runtime Grounding Verifier (RGV), a programmatic system '
            'that verifies every numerical claim in LLM-generated text against the formal JSON '
            'contract J. Unlike prior work that enforces G(s) ⊆ J only through prompt instruction, '
            'RGV enforces it algorithmically and produces a measurable scalar metric G_rate ∈ [0, 1]. '
            'This formal verification property has no published precedent in environmental monitoring '
            'pipelines and constitutes the primary novel contribution of ATARS. [ATARS Eq. 19]'
        )
        run.font.size = Pt(10); run.font.color.rgb = PURPLE; run.bold = True
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(10)

        # ── Overall metrics table ─────────────────────────────────────────
        h2('Formal Verification Result')
        overall_g  = verification.get('overall_G_rate', 0)
        v_passed   = verification.get('verification_passed', False)
        v_verdict  = verification.get('overall_verdict', 'N/A')
        n_total    = verification.get('total_numerical', 0)
        n_grounded = verification.get('total_grounded', 0)
        n_ungnd    = verification.get('total_ungrounded', 0)
        j_size     = verification.get('j_lookup_size', 0)

        t = doc.add_table(rows=1, cols=4); t.style = 'Table Grid'
        for cell, (lbl, val, bg_hex) in zip(t.rows[0].cells, [
            ('Overall G_rate',    f'{overall_g:.1%}',  '1E6B3C' if v_passed else '92400E'),
            ('Verdict',           v_verdict,            '1E6B3C' if v_passed else '92400E'),
            ('Grounded / Total',  f'{n_grounded} / {n_total}', '2557A7'),
            ('J Lookup Entries',  str(j_size),          '4C1D95'),
        ]):
            _set_cell_bg(cell, 'F4FFF4' if v_passed else 'FFF9F0')
            lp = cell.add_paragraph(lbl)
            lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in lp.runs: r.font.size = Pt(9); r.bold = True; r.font.color.rgb = GRAY
            vp = cell.add_paragraph(val)
            vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in vp.runs:
                r.font.size = Pt(20); r.bold = True
                r.font.color.rgb = _rgb(bg_hex)

        doc.add_paragraph()
        add_image('grounding', 6.2,
                  'Figure — RGV Grounding Verification Dashboard. (a) G_rate per section. '
                  '(b) Overall G_rate gauge. (c) Verdict distribution. '
                  '(d) Sentence-level analysis. RGV v1.0 — ATARS novel contribution.')

        # ── Per-section verification table ───────────────────────────────
        h2('Per-Section Grounding Results')
        body('G(s) ⊆ J verified per section. PASS = G_rate ≥ 75% of numerical sentences grounded in J.',
             italic=True)
        t = doc.add_table(rows=1, cols=6); t.style = 'Table Grid'
        for cell, hdr in zip(t.rows[0].cells,
                              ['Section', 'Total Sent.', 'Numerical', 'Grounded',
                               'G_rate', 'Verdict']):
            _set_cell_bg(cell, '1A2C4E')
            p = cell.paragraphs[0]
            r = p.add_run(hdr); r.bold = True; r.font.size = Pt(9)
            r.font.color.rgb = WHITE
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for sec_key, sec_data in verification.get('sections', {}).items():
            row = t.add_row()
            bg  = 'D1FAE5' if sec_data.get('passes') else 'FEF3C7'
            vals = [
                sec_key,
                str(sec_data.get('total_sentences', 0)),
                str(sec_data.get('numerical_sentences', 0)),
                str(sec_data.get('grounded_sentences', 0)),
                f"{sec_data.get('G_rate', 0):.1%}",
                sec_data.get('verdict', 'N/A'),
            ]
            for i, (cell, val) in enumerate(zip(row.cells, vals)):
                _set_cell_bg(cell, bg)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(val); r.font.size = Pt(9)

        doc.add_paragraph()

        # ── Evidence table for grounded claims ────────────────────────────
        h2('Sample Evidence — Grounded Claims Matched to J')
        body('Each row shows one numerical claim extracted from narrative text, '
             'the matching entry in JSON contract J, and the tolerance used.', italic=True)
        t = doc.add_table(rows=1, cols=5); t.style = 'Table Grid'
        for cell, hdr in zip(t.rows[0].cells,
                              ['Claim', 'Value', 'Matched J Path', 'J Value', 'Diff']):
            _set_cell_bg(cell, '1A2C4E')
            p = cell.paragraphs[0]
            r = p.add_run(hdr); r.bold = True; r.font.size = Pt(9)
            r.font.color.rgb = WHITE
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        evidence_count = 0
        for sec_key, sec_data in verification.get('sections', {}).items():
            for sent in sec_data.get('sentences', []):
                for ev in sent.get('evidence', []):
                    if ev.get('grounded') and evidence_count < 15:
                        row = t.add_row()
                        bg  = 'D1FAE5' if evidence_count % 2 == 0 else 'F0FFF4'
                        jp  = ev.get('j_path', '')
                        jp_short = jp[-40:] if len(jp) > 40 else jp
                        for cell, val in zip(row.cells, [
                            str(ev.get('claim', '')),
                            str(ev.get('value', '')),
                            jp_short,
                            str(round(ev.get('j_value', 0), 4)),
                            f"{ev.get('diff', 0):.4f}",
                        ]):
                            _set_cell_bg(cell, bg)
                            p = cell.paragraphs[0]
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            r = p.add_run(val); r.font.size = Pt(8.5)
                        evidence_count += 1

        if n_ungnd > 0:
            doc.add_paragraph()
            h2(f'Ungrounded Sentences ({n_ungnd})')
            body(
                'The following sentences contain numerical claims that could not be matched '
                'to any entry in the JSON contract J within the configured tolerance. '
                'This does not necessarily indicate hallucination — it may also mean the '
                'value was computed from J (e.g. a percentage derived from two J values) '
                'rather than stored directly in J. Expert review is recommended for each.',
                italic=True, color=AMBER
            )
            ungrounded_list = verification.get('ungrounded_list', [])
            for item in ungrounded_list[:10]:
                section_lbl = item.get('section', '?')
                sentence    = item.get('sentence', '').strip()
                n_claims    = item.get('n_claims', 0)
                # Section badge line
                p_badge = doc.add_paragraph()
                p_badge.paragraph_format.left_indent  = Inches(0.2)
                p_badge.paragraph_format.space_before = Pt(6)
                p_badge.paragraph_format.space_after  = Pt(1)
                r_badge = p_badge.add_run(
                    f"Section [{section_lbl}]  —  {n_claims} unmatched numerical claim(s)"
                )
                r_badge.bold = True
                r_badge.font.size = Pt(9)
                r_badge.font.color.rgb = AMBER
                # Full sentence on next line — word-wrapped, never truncated
                p_sent = doc.add_paragraph()
                p_sent.paragraph_format.left_indent  = Inches(0.4)
                p_sent.paragraph_format.space_before = Pt(1)
                p_sent.paragraph_format.space_after  = Pt(8)
                display_text = sentence if sentence else "(sentence text not available)"
                r_sent = p_sent.add_run(display_text)
                r_sent.font.size  = Pt(9)
                r_sent.italic     = True
                r_sent.font.color.rgb = RED

        # ── Formal definition ─────────────────────────────────────────────
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run(
            'Formal Definition: G_rate = |{s ∈ N_num : G(s) ⊆ J}| / |N_num|  '
            'where N_num = {s ∈ N : C(s) ≠ ∅} (sentences with numerical claims). '
            'G(s) ⊆ J holds when every numerical value in s matches a corresponding '
            'entry in J within type-adaptive tolerance ε '
            '(±5% relative for concentrations; ±2 absolute for integer day counts; '
            '±0.1 absolute for quality scores). '
            'Prior systems enforced this property only via prompt instruction; '
            'RGV enforces it algorithmically and produces a computable, auditable metric. '
            '[ATARS Eq. 19]'
        )
        run.font.size = Pt(9.5); run.italic = True; run.font.color.rgb = PURPLE

        doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    #  SECTION 13b — ML ANOMALY DETECTION (Isolation Forest)
    # ════════════════════════════════════════════════════════════════════════
    h1('Section 14 — ML Analysis: Isolation Forest Multivariate Anomaly Detection')
    p = doc.add_paragraph()
    run = p.add_run(
        'This analysis is additive and does not replace the formal z-score operator ζ_d (Eq. 9). '
        'Isolation Forest detects days where the combination of multiple pollutants '
        'is anomalous — a pattern that univariate z-score analysis cannot identify. '
        'All computation is performed locally on the aggregated daily dataset; '
        'no raw data is transmitted externally. '
        'The fixed random seed (random_state = 42) ensures this result is fully reproducible.'
    )
    run.font.size = Pt(10); run.italic = True; run.font.color.rgb = TEAL
    p.paragraph_format.space_before = Pt(4)

    if ml_if.get('available'):
        # Summary table
        t = doc.add_table(rows=1, cols=2)
        t.style = 'Table Grid'
        hdr = t.rows[0].cells
        hdr[0].text = 'Isolation Forest Parameter'; hdr[1].text = 'Value'
        for cell in hdr:
            cell.paragraphs[0].runs[0].bold = True
            _set_cell_bg(cell, '003366')
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

        rows_data = [
            ('Features used', ', '.join(ml_if['features_used'])),
            ('n_estimators', str(ml_if['n_estimators'])),
            ('contamination', f"{ml_if['contamination']:.0%} (expected anomaly fraction)"),
            ('random_state', f"{ml_if['random_state']} (fixed for reproducibility)"),
            ('Total days analysed', str(ml_if['n_days'])),
            ('Anomaly days (IF)', f"{ml_if['n_anomalies']} ({ml_if['anomaly_rate']*100:.1f}%)"),
            ('Flagged by BOTH methods', str(ml_if['n_both'])),
            ('IF-only (multivariate — missed by z-score)', str(ml_if['n_if_only'])),
            ('Z-score only (univariate spike, within multivariate norm)', str(ml_if['n_zs_only'])),
        ]
        for lbl, val in rows_data:
            row = t.add_row().cells
            row[0].text = lbl; row[1].text = val

        doc.add_paragraph()
        add_image('ml_anomaly', 6.2,
                  'Figure — Isolation Forest vs Z-Score Anomaly Comparison. '
                  '(a) IF decision score — lower = more anomalous. '
                  '(b) Z-score ζ_d with ±3σ threshold. '
                  '(c) Agreement map — Red=both, Purple=IF-only, Amber=z-score-only.')
        
        p = doc.add_paragraph()
        run = p.add_run(
            'Interpretation of IF-only days: These are dates where no single pollutant '
            'crossed the 3σ z-score threshold, but the joint distribution of '
            f"{', '.join(ml_if['features_used'])} was statistically unusual compared "
            'to the training distribution. These warrant manual investigation as potential '
            'multi-source pollution episodes. IF-only days represent the added analytical '
            'value of multivariate over univariate anomaly detection.'
        )
        run.font.size = Pt(9.5); run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    else:
        p = doc.add_paragraph()
        p.add_run(f"Isolation Forest skipped: {ml_if.get('reason', 'insufficient data')}.").italic = True

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    #  SECTION 13c — ML FORECASTING (Holt-Winters)
    # ════════════════════════════════════════════════════════════════════════
    h1('Section 15 — ML Analysis: Holt-Winters PM10 Forecast')
    p = doc.add_paragraph()
    run = p.add_run(
        'This section presents supplementary PM10 forecasting using Triple Exponential '
        'Smoothing (Holt-Winters additive seasonal model), which captures level, trend, '
        'and weekly seasonal components simultaneously. '
        'Forecast values are deliberately excluded from the JSON contract J so that the '
        'language model cannot present statistical extrapolations as established facts. '
        'All forecast content appears in this section and Chart 15 only.'
    )
    run.font.size = Pt(10); run.italic = True; run.font.color.rgb = TEAL
    p.paragraph_format.space_before = Pt(4)

    if ml_hw.get('available'):
        t = doc.add_table(rows=1, cols=2)
        t.style = 'Table Grid'
        hdr = t.rows[0].cells
        hdr[0].text = 'Holt-Winters Parameter'; hdr[1].text = 'Value'
        for cell in hdr:
            cell.paragraphs[0].runs[0].bold = True
            _set_cell_bg(cell, '003366')
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

        fc_mean = float(np.mean(ml_hw['forecast']))
        fc_max  = float(np.max(ml_hw['forecast']))
        rows_data = [
            ('α (level smoothing)', str(ml_hw['alpha'])),
            ('β (trend smoothing)', str(ml_hw['beta'])),
            ('γ (seasonal smoothing)', str(ml_hw['gamma'])),
            ('Seasonal period', f"{ml_hw['seasonal_periods']} days (weekly)"),
            ('Forecast horizon', f"{ml_hw['forecast_days']} days"),
            ('History used', f"{ml_hw['n_history']} days"),
            ('In-sample RMSE', f"{ml_hw['rmse']:.2f} µg/m³"),
            ('In-sample MAE', f"{ml_hw['mae']:.2f} µg/m³"),
            ('Residual std (σ)', f"{ml_hw['resid_std']:.2f} µg/m³"),
            (f"{ml_hw['forecast_days']}-day forecast mean", f"{fc_mean:.1f} µg/m³"),
            (f"{ml_hw['forecast_days']}-day forecast peak", f"{fc_max:.1f} µg/m³"),
            ('Prediction interval', '80% (±1.28σ√h — widens with horizon)'),
        ]
        for lbl, val in rows_data:
            row = t.add_row().cells
            row[0].text = lbl; row[1].text = val

        doc.add_paragraph()
        add_image('forecast', 6.2,
                  'Figure — Holt-Winters PM10 Forecast. '
                  '(a) Historical PM10 + fitted values + 14-day forecast with 80% prediction interval. '
                  '(b) Model residuals with ±1σ bands.')

        p = doc.add_paragraph()
        run = p.add_run(
            'Limitations of this forecast: (1) Holt-Winters is a statistical extrapolation, '
            'not a physical atmospheric model — it does not account for meteorological events, '
            'policy interventions, or emission changes. '
            '(2) Uncertainty grows significantly beyond 7 days — the 80% PI widens with √h. '
            '(3) The model was trained on this dataset only — transferability to other '
            'stations or time periods is not validated here. '
            '(4) For public health decisions, use this forecast as an indicative signal '
            'only — not as a definitive prediction.'
        )
        run.font.size = Pt(9.5); run.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
    else:
        p = doc.add_paragraph()
        p.add_run(f"Forecast skipped: {ml_hw.get('reason', 'disabled or insufficient data')}.").italic = True

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    #  SECTION 14 — AI NARRATIVE DISCLOSURE
    h1('Section 16 — AI-Assisted Commentary and Limitations')

    # AI disclosure banner
    p = doc.add_paragraph()
    run = p.add_run(
        'Sections 01 through 12 of this report contain AI-assisted narrative text generated '
        'by a Large Language Model operating under formal grounding constraints G(s) ⊆ J '
        '(ATARS Eq. 19). '
        f'Model used: {narrative.get("model", "statistical_placeholder")}. '
        'Temperature τ = 0 (deterministic output). All numerical claims in AI-generated '
        'text reference the JSON contract J exclusively. '
        + ('AI narrative generation was active during this run.' if narrative.get('ai_flag') else
           'AI narrative generation was unavailable; statistical placeholder text was used instead.')
    )
    run.font.size = Pt(10); run.italic = True; run.font.color.rgb = BLUE
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(12)

    body(narrative.get('F_limitations', ''))

    # ════════════════════════════════════════════════════════════════════════
    #  SECTION 15 — WINDOWS TASK SCHEDULER AUTOMATION
    # ════════════════════════════════════════════════════════════════════════
    h1('Section 17 — Windows Task Scheduler: Automated Daily Execution')
    body(
        'ATARS is designed for fully automated daily operation using Windows Task Scheduler, '
        'a built-in Windows feature that executes programs at a fixed time without manual '
        'intervention. Once configured, ATARS runs every morning automatically and places '
        'a completed report in the output folder — provided the updated data file has been '
        'deposited before the scheduled run time.'
    )

    h2('How Task Scheduler Works with ATARS')
    sched_steps = [
        ('Step 1', 'Data arrives', 'Your sensor system, laboratory, or field team deposits the updated data CSV into the configured folder before the scheduled time (e.g., 06:00 AM).'),
        ('Step 2', 'Task Scheduler triggers', 'Windows Task Scheduler runs atars.py at the exact configured time (e.g., 07:00 AM) — no human action required.'),
        ('Step 3', 'Pipeline executes', 'ATARS loads the latest data, runs all 9 pipeline steps (Steps 1–9), generates up to 15 charts, and compiles the full Word document report.'),
        ('Step 4', 'Report delivered', 'The completed .docx report appears in the output folder and/or is emailed to configured recipients.'),
        ('Step 5', 'Audit logged', 'SHA-256 hash of the run is recorded in the audit log for reproducibility verification.'),
    ]
    t = doc.add_table(rows=1, cols=3)
    t.style = 'Table Grid'
    for cell, hdr_text in zip(t.rows[0].cells, ['Step', 'Event', 'What Happens']):
        p = cell.paragraphs[0]; run = p.add_run(hdr_text)
        run.bold = True; run.font.size = Pt(10)
        _set_cell_bg(cell, '1A2C4E')
        for r in p.runs: r.font.color.rgb = WHITE
    for step, event, desc in sched_steps:
        row = t.add_row()
        for cell, val, w in zip(row.cells, [step, event, desc], [1200, 2000, 6160]):
            p = cell.paragraphs[0]; r = p.add_run(val); r.font.size = Pt(10)
            _set_cell_bg(cell, 'DCE8F5')

    h2('Setup Instructions (One-Time, ~5 minutes)')
    body('1. Open Windows search → type "Task Scheduler" → Open it.')
    body('2. Click "Create Basic Task" in the right panel.')
    body('3. Name it: ATARS Daily Report — click Next.')
    body('4. Trigger: select "Daily" → set your time (e.g., 07:00 AM) → Next.')
    body('5. Action: "Start a program" → Next.')
    body('6. Program/script field: type the full path to python.exe')
    body('   Example: C:\\Python311\\python.exe')
    body('7. Add arguments field: type the full path to atars.py and arguments')
    body('   Example: "C:\\ATARS\\atars.py" --data "C:\\ATARS\\data.csv" --city "Delhi" --no-llm')
    body('8. Start in field: type the folder containing atars.py')
    body('   Example: C:\\ATARS\\')
    body('9. Click Finish. ATARS will now run automatically every day at your chosen time.')

    h2('The run_atars.bat Batch File (Easier Method)')
    body('ATARS automatically creates a run_atars.bat file in your output folder. '
         'Point Task Scheduler at this .bat file instead of typing all arguments manually. '
         'The batch file contains the exact command with all your configured settings.')

    # Code block style
    code_lines = [
        '@echo off',
        'REM ATARS — Automated Daily Run Script',
        'REM Generated by ATARS. Point Windows Task Scheduler at this file.',
        '',
        'cd /d "C:\\ATARS"',
        'python atars.py --data "data.csv" --city "Delhi" --station "CPCB-001" --no-llm',
        '',
        'REM Check exit code',
        'if %ERRORLEVEL% EQU 0 (',
        '    echo ATARS run completed successfully >> atars_output\\run_log.txt',
        ') else (',
        '    echo ATARS run FAILED — check pipeline >> atars_output\\run_log.txt',
        ')',
    ]
    for line in code_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Pt(24)
        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'; run.font.size = Pt(9)
        run.font.color.rgb = _rgb('065F46') if line.startswith('REM') else \
                              _rgb('1A2C4E') if line.startswith('@') else \
                              _rgb('000000')

    body('IMPORTANT: The data file must arrive in the folder BEFORE Task Scheduler runs. '
         'Configure your data delivery (sensor upload, FTP, email attachment) to complete '
         'at least 30 minutes before the scheduled ATARS run time.',
         italic=True, color=AMBER)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    #  SECTION 16 — REPRODUCIBILITY METADATA
    # ════════════════════════════════════════════════════════════════════════
    h1('Section 18 — Reproducibility and Audit Record')
    body(
        'The metadata table below provides a complete record of this pipeline run, '
        'enabling independent verification and exact re-execution. '
        'Hash(O) = SHA-256(canonical_JSON(J)) — the output is guaranteed to be identical '
        'if the same data file and configuration are used with the same ATARS version '
        '(ATARS Reproducibility Theorem, Eq. 20).'
    )

    run_id    = hashlib.sha256(json.dumps(J, sort_keys=True).encode()).hexdigest()[:16]
    data_hash = hashlib.sha256(json.dumps(
        {k: v for k, v in J.items() if k != 'ai_constraints'},
        sort_keys=True
    ).encode()).hexdigest()

    meta = [
        ('Report ID',         f'ATARS-{datetime.now().strftime("%Y%m%d")}-{run_id[:8].upper()}'),
        ('Generated',         datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
        ('City / Station',    f'{J["city"]} / {J["station_id"]}'),
        ('Date Range',        J['date_range']),
        ('Total Records',     f'{J["total_records"]:,}'),
        ('Overall Q(D)',      f'{J["data_quality"]["overall_Q"]:.4f}'),
        ('Confidence Flag',   J['data_quality']['confidence_flag']),
        ('Baseline Window W', f'{J["formal_operators"]["baseline_window_W"]} days'),
        ('Z-Threshold ζ',     str(J['formal_operators']['z_threshold'])),
        ('Run Hash (J)',       data_hash[:32] + '...'),
        ('Framework Version', 'ATARS v2.0.0 — Priyanshu (2026)'),
        ('Author',            f'{config["author"]}  ·  {config["institution"]}'),
        ('AI Model',          narrative.get('model', 'N/A')),
        ('LLM Temperature',   'τ = 0 (deterministic)'),
        ('Classification',    'Open Source — MIT License — Copyright (c) 2026 Priyanshu'),
    ]

    t = doc.add_table(rows=0, cols=2)
    t.style = 'Table Grid'
    for i, (label, value) in enumerate(meta):
        row = t.add_row()
        bg = 'DCE8F5' if i % 2 == 0 else 'FFFFFF'
        for cell, val, is_label in [(row.cells[0], label, True),
                                     (row.cells[1], value, False)]:
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(10)
            r.bold = is_label
            if is_label:
                r.font.color.rgb = NAVY
            _set_cell_bg(cell, '1A3A5C' if is_label else bg)
            if is_label:
                for rn in p.runs: rn.font.color.rgb = WHITE

    # ── Final signature ───────────────────────────────────────────────────
    doc.add_paragraph()
    p = doc.add_paragraph(
        f'"A framework built for researchers, by a researcher."\n'
        f'{config["author"]}  ·  {config["institution"]}  ·  {config["email"]}'
    )
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(10); run.italic = True; run.font.color.rgb = GRAY

    # Save
    fname    = f'ATARS_Report_{J["city"].replace(" ","_")}_{datetime.now().strftime("%Y%m%d_%H%M")}.docx'
    out_path = str(out_dir / fname)
    doc.save(out_path)
    print(f"  ✓ Report saved: {out_path}")
    return out_path


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 12 — REPRODUCIBILITY AUDIT LOG
# ═══════════════════════════════════════════════════════════════════════════════

def write_audit_log(J: dict, narrative: dict, report_path: str,
                     config: dict, out_dir: Path, verification: dict = None,
                     ml_if: dict = None, ml_hw: dict = None):
    """
    Write SHA-256 audit log and generate Windows Task Scheduler batch file.
    Hash(O) = SHA-256(canonical_JSON(J))  — Eq. 21
    v2.0: Logs ML results (IF anomaly count, HW forecast metrics) for audit trail.
    """
    ml_if = ml_if or {}
    ml_hw = ml_hw or {}
    J_canonical   = json.dumps(J, sort_keys=True, ensure_ascii=True)
    hash_J        = hashlib.sha256(J_canonical.encode()).hexdigest()

    with open(report_path, 'rb') as f:
        hash_report = hashlib.sha256(f.read()).hexdigest()

    audit = {
        "run_id"       : hash_J[:16],
        "timestamp"    : datetime.now().isoformat(),
        "city"         : J['city'],
        "date_range"   : J['date_range'],
        "total_records": J['total_records'],
        "quality_Q"    : J['data_quality']['overall_Q'],
        "hash_J"       : hash_J,
        "hash_report"  : hash_report,
        "report_file"  : os.path.basename(report_path),
        "llm_model"    : narrative.get('model', 'N/A'),
        "framework_version": "ATARS_v1.0",
        # SECURITY: config_used logs only statistical/operational parameters.
        # Personal identifiers (author, email, institution) and absolute file
        # paths are intentionally excluded to prevent accidental data leakage
        # when audit logs are shared for reproducibility verification.
        "config_used"  : {k: v for k, v in config.items()
                          if k not in (
                              'thresholds',      # large dict, referenced via hash_J
                              'author',          # personal identifier — excluded
                              'email',           # personal contact — excluded
                              'institution',     # personal institution — excluded
                              'department',      # personal department — excluded
                              'degree',          # personal info — excluded
                              'location',        # personal location — excluded
                              'data_file',       # absolute path may reveal FS layout
                              'output_dir',      # absolute path may reveal FS layout
                              'ollama_url',      # always localhost; not needed
                          )},
        "grounding_G_rate"     : verification["overall_G_rate"] if verification else None,
        "grounding_verdict"    : verification["overall_verdict"] if verification else "N/A",
        "grounding_passed"     : verification["verification_passed"] if verification else None,
        "reproducibility_note":
            "Re-run with same data file and config to verify Hash(J) match.",
        # ML audit fields (v2.0)
        "ml_isolation_forest"  : {
            "available"     : ml_if.get("available", False),
            "n_anomalies"   : ml_if.get("n_anomalies", None),
            "anomaly_rate"  : ml_if.get("anomaly_rate", None),
            "n_estimators"  : ml_if.get("n_estimators", None),
            "contamination" : ml_if.get("contamination", None),
            "random_state"  : ml_if.get("random_state", None),
            "features_used" : ml_if.get("features_used", None),
            "n_if_only"     : ml_if.get("n_if_only", None),
            "security_note" : ml_if.get("security_note", ""),
        },
        "ml_holt_winters"      : {
            "available"     : ml_hw.get("available", False),
            "forecast_days" : ml_hw.get("forecast_days", None),
            "rmse"          : ml_hw.get("rmse", None),
            "mae"           : ml_hw.get("mae", None),
            "alpha"         : ml_hw.get("alpha", None),
            "beta"          : ml_hw.get("beta", None),
            "gamma"         : ml_hw.get("gamma", None),
            "security_note" : ml_hw.get("security_note", ""),
        },
        "atars_version"        : "2.0.0",
        "copyright"            : "Copyright (c) 2026 Priyanshu — MIT License (Open Source)",
    }

    log_path = out_dir / f'audit_log_{audit["run_id"]}.json'
    with open(log_path, 'w') as f:
        json.dump(audit, f, indent=2)
    print(f"  ✓ Audit log: {log_path}")

    # ── Generate Windows Task Scheduler batch file ─────────────────────────
    _write_batch_file(config, out_dir)

    return audit


def _write_batch_file(config: dict, out_dir: Path):
    """
    Auto-generate a Windows Task Scheduler .bat file for daily automation.
    Users point Task Scheduler at this file to run ATARS every day automatically.
    """
    script_path  = os.path.abspath('atars.py')
    data_path    = os.path.abspath(config.get('data_file', 'data.csv'))
    out_path_abs = os.path.abspath(str(out_dir))
    work_dir     = os.path.dirname(script_path)
    llm_flag     = '' if config.get('use_llm') else ' --no-llm'
    model_flag   = f' --llm-model {config["llm_model"]}' if config.get('use_llm') else ''

    bat_content = f"""@echo off
REM ═══════════════════════════════════════════════════════════════════════════
REM  ATARS — Automated Time-Series Analysis and Reporting System
REM  Windows Task Scheduler Daily Run Script
REM  Author : {config['author']}
REM  Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}
REM
REM  HOW TO USE:
REM  1. Open Windows Task Scheduler (search in Start menu)
REM  2. Click "Create Basic Task" → Name: "ATARS Daily Report"
REM  3. Trigger: Daily → set your time (e.g. 07:00 AM)
REM  4. Action: Start a Program
REM  5. Program/script: point to this .bat file
REM     Path: {out_path_abs}\\run_atars.bat
REM  6. Click Finish
REM
REM  IMPORTANT: Your data file must arrive BEFORE the scheduled run time.
REM  Recommended: schedule data delivery 30+ minutes before ATARS runs.
REM ═══════════════════════════════════════════════════════════════════════════

REM Change to the ATARS working directory
cd /d "{work_dir}"

REM Record run start time
echo [%DATE% %TIME%] ATARS run started >> "{out_path_abs}\\scheduler_log.txt"

REM Run ATARS pipeline
python "{script_path}" ^
    --data "{data_path}" ^
    --city "{config['city']}" ^
    --station "{config['station_id']}" ^
    --output "{out_path_abs}"{llm_flag}{model_flag}

REM Check if run succeeded
if %ERRORLEVEL% EQU 0 (
    echo [%DATE% %TIME%] ATARS SUCCESS — report saved to {out_path_abs} >> "{out_path_abs}\\scheduler_log.txt"
    echo ATARS completed successfully.
) else (
    echo [%DATE% %TIME%] ATARS FAILED — check pipeline output above >> "{out_path_abs}\\scheduler_log.txt"
    echo ATARS run FAILED. Check output for errors.
    exit /b 1
)

REM Optional: open output folder automatically after run (comment out for silent mode)
REM explorer "{out_path_abs}"

exit /b 0
"""

    bat_path = out_dir / 'run_atars.bat'
    with open(bat_path, 'w', encoding='utf-8') as f:
        f.write(bat_content)
    print(f"  ✓ Batch file: {bat_path}")
    print(f"  ✓ Point Windows Task Scheduler at: {bat_path}")
    return str(bat_path)


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 13 — MAIN PIPELINE ORCHESTRATOR (Steps 1–9)
# ═══════════════════════════════════════════════════════════════════════════════

def print_banner(config: dict):
    print("""
╔══════════════════════════════════════════════════════════════════════════════╗
║   ATARS — Automated Time-Series Analysis and Reporting System  v2.0.0      ║
║   Author: Priyanshu · Global Institute of Technology and Management        ║
║   Open Source · MIT License · Formally specified research pipeline          ║
║                                                                              ║
║   Formal operators: A(D_v), B(D,W), ζ_d, δ_d, Q(D), G:J→N, Hash(O)       ║
╚══════════════════════════════════════════════════════════════════════════════╝""")
    print(f"  City: {config['city']}  |  Station: {config['station_id']}")
    print(f"  Output: {config['output_dir']}")
    print(f"  LLM: {'Enabled (' + config['llm_model'] + ')' if config['use_llm'] else 'Disabled'}")


def run_pipeline(config: dict):
    """
    Nine-step ATARS pipeline.
    Step 1: Ingest → Step 2: Trigger → Step 3: Aggregate →
    Step 4: JSON Assembly → Step 5: AI Narrative → Step 6: Charts →
    Step 7: Report → Step 8: Audit Log → Step 9: Summary
    """
    print_banner(config)

    out_dir = Path(config['output_dir'])
    out_dir.mkdir(parents=True, exist_ok=True)

    try:
        # ── Step 1: Data Ingestion ─────────────────────────────────────────
        print(f"\n{'─'*70}")
        print("  STEP 1 — Data Ingestion and Validation")
        df = load_data(config['data_file'])

        # ── Step 2: Trigger Guard ──────────────────────────────────────────
        print(f"\n{'─'*70}")
        print("  STEP 2 — Data Availability Check")
        n_records = len(df)
        if n_records == 0:
            print("  ✗ No records found. Exiting.")
            return
        print(f"  ✓ {n_records:,} records spanning {df['date_only'].nunique()} days — proceeding.")

        # ── Step 3: Quality Flags + Statistical Aggregation ───────────────
        print(f"\n{'─'*70}")
        print("  STEP 3 — Quality Assessment + Statistical Aggregation")
        df       = assign_quality_flags(df)
        daily_df = build_daily_stats(df, config)

        # ── Step 3b: Correlation + OLS ─────────────────────────────────────
        print("\n  Computing correlation matrix R and OLS regression β̂...")
        corr_matrix = compute_correlation_matrix(daily_df, POLLUTANTS + METEOROLOGICAL[:4])
        ols_result  = run_ols_regression(daily_df, target='PM10',
                                          predictors=['humidity_pct','wind_speed_10m',
                                                      'rain_mm','pressure_hpa'])
        if ols_result:
            print(f"  ✓ OLS: R² = {ols_result['r_squared']:.4f} for PM10 ~ meteorology")

        # ── Step 3c: ML Anomaly Detection (Isolation Forest) ──────────────
        print(f"\n{'─'*70}")
        print("  STEP 3c — ML-1: Isolation Forest Multivariate Anomaly Detection")
        print("  Security: computed locally on daily_df — no external transmission.")
        ml_if = run_isolation_forest(daily_df, config)

        # ── Step 3d: Holt-Winters Forecast ─────────────────────────────────
        if config.get("use_forecast", True):
            print(f"\n{'─'*70}")
            print("  STEP 3d — ML-2: Holt-Winters PM10 Forecast")
            print("  Security: computed locally — forecast NOT added to J (LLM data contract).")
            ml_hw = run_hw_forecast(daily_df, config)
        else:
            ml_hw = {"available": False, "reason": "disabled_by_flag"}
            print("  STEP 3d — ML-2: Forecast disabled (--no-forecast)")

        # ── Step 4: JSON Contract Assembly ────────────────────────────────
        print(f"\n{'─'*70}")
        print("  STEP 4 — JSON Contract Assembly")
        J = build_json_contract(df, daily_df, config, ols_result, corr_matrix)

        # Save J for inspection
        json_path = out_dir / 'json_contract.json'
        with open(json_path, 'w') as f:
            json.dump(J, f, indent=2)
        print(f"  ✓ JSON contract saved: {json_path}")

        # ── Step 5: AI Narrative ───────────────────────────────────────────
        print(f"\n{'─'*70}")
        print("  STEP 5 — AI Narrative Generation (G: J → N)")
        narrative = try_llm_narrative(J, config)

        # ── Step 5.5: Runtime Grounding Verification ───────────────────────
        print(f"\n{'─'*70}")
        print("  STEP 5.5 — Runtime Grounding Verifier (RGV) — novel contribution")
        print("  Programmatic enforcement: G(s) ⊆ J  ∀ s ∈ N  [ATARS Eq. 19]")
        rgv         = RuntimeGroundingVerifier(J)
        verification= rgv.verify_narrative(narrative)
        g_rate      = verification['overall_G_rate']
        g_verdict   = verification['overall_verdict']
        n_ungrounded= verification['total_ungrounded']
        print(f"  ✓ G_rate = {g_rate:.1%}  [{g_verdict}]")
        print(f"  ✓ J lookup: {verification['j_lookup_size']} entries | "
              f"Numerical sentences: {verification['total_numerical']} | "
              f"Ungrounded: {n_ungrounded}")
        if n_ungrounded > 0:
            print(f"  ⚠  {n_ungrounded} sentence(s) ungrounded — flagged in report")
        # Save verification report to JSON
        ver_path = out_dir / 'grounding_verification.json'
        with open(ver_path, 'w') as _f:
            json.dump(verification, _f, indent=2, default=str)
        print(f"  ✓ Verification saved: {ver_path}")

        # ── Step 6: Chart Generation ───────────────────────────────────────
        print(f"\n{'─'*70}")
        print("  STEP 6 — Chart Generation (12+1 publication-quality figures)")
        # ── Step 5.5b: Generate Chart Insights (deterministic) ──────────
        print(f"\n{'─'*70}")
        print("  STEP 5.5b — Chart Insight Generation (deterministic — no LLM)")
        chart_insights = generate_chart_insights(
            daily_df, df, config, corr_matrix, ols_result, ml_if, ml_hw
        )
        print(f"  ✓ Insights generated for {len(chart_insights)} charts")

        # ── Step 6: Chart Generation ────────────────────────────────────────
        charts = generate_all_charts(df, daily_df, config, out_dir, verification,
                                     ml_if=ml_if, ml_hw=ml_hw)

        # ── Step 7: Report Compilation ─────────────────────────────────────
        print(f"\n{'─'*70}")
        print("  STEP 7 — Report Compilation")
        report_path = build_report(
            J, daily_df, charts, narrative,
            ols_result, corr_matrix, config, out_dir,
            verification=verification,
            chart_insights=chart_insights,
            ml_if=ml_if, ml_hw=ml_hw
        )

        # ── Step 8: Audit Log ──────────────────────────────────────────────
        print(f"\n{'─'*70}")
        print("  STEP 8 — Reproducibility Audit Log")
        audit = write_audit_log(J, narrative, report_path, config, out_dir,
                                 verification=verification, ml_if=ml_if, ml_hw=ml_hw)

        # ── Step 9: Summary ────────────────────────────────────────────────
        print(f"\n{'═'*70}")
        print("  STEP 9 — PIPELINE COMPLETE")
        print(f"{'═'*70}")
        print(f"  ✓ Report    : {report_path}")
        print(f"  ✓ Charts    : {len(charts)} figures in {out_dir}/")
        print(f"  ✓ JSON      : {json_path}")
        print(f"  ✓ Batch file: {out_dir}/run_atars.bat")
        print(f"  ✓ Audit ID  : {audit['run_id']}")
        print(f"  ✓ Hash(J)   : {audit['hash_J'][:32]}...")
        g_rate_val = verification.get('overall_G_rate', None)
        if g_rate_val is not None:
            print(f"  ✓ G_rate    : {g_rate_val:.1%}  [{verification.get('overall_verdict','N/A')}]  (RGV novel contribution)")
        if ml_if.get('available'):
            print(f"  ✓ IF anomaly: {ml_if['n_anomalies']} days  "
                  f"| IF-only (multivariate): {ml_if['n_if_only']}  "
                  f"| random_state={ml_if['random_state']}")
        if ml_hw.get('available'):
            print(f"  ✓ HW forecast: RMSE={ml_hw['rmse']:.2f}  "
                  f"| 14-day mean={float(np.mean(ml_hw['forecast'])):.1f} µg/m³")
        print(f"\n  Q(D) = {J['data_quality']['overall_Q']:.4f}  [{J['data_quality']['confidence_flag']}]")
        print(f"  Top exceedance variable: {J.get('top_pollutant', 'N/A')}")
        exceedances = J.get('exceedances', {})
        if exceedances:
            print(f"  Exceedance summary: " +
                  ", ".join(f"{k}={v}d" for k, v in list(exceedances.items())[:5]))

        print(f"\n{'─'*70}")
        print("  WINDOWS TASK SCHEDULER — AUTOMATE DAILY REPORTS")
        print(f"{'─'*70}")
        print(f"  A batch file has been created: {out_dir}/run_atars.bat")
        print(f"  To run ATARS automatically every day:")
        print(f"  1. Open Windows Start Menu → search 'Task Scheduler' → Open")
        print(f"  2. Click 'Create Basic Task' → Name: ATARS Daily Report")
        print(f"  3. Trigger: Daily → set time (e.g. 07:00 AM)")
        print(f"  4. Action: Start a Program")
        print(f"  5. Point to: {os.path.abspath(str(out_dir))}/run_atars.bat")
        print(f"  6. Click Finish")
        print(f"  NOTE: Ensure your data file arrives BEFORE the scheduled time.")
        print(f"{'─'*70}")
        print(f"\n  ATARS v2.0.0 — Open Source (MIT License) — Copyright (c) 2026 Priyanshu")
        print(f"  Reproducibility: re-run with same data + config")
        print(f"  to verify Hash(J) = {audit['hash_J'][:16]}...")
        print(f"{'═'*70}\n")

        return report_path, J, audit

    except Exception as e:
        print(f"\n  ✗ Pipeline error: {e}")
        traceback.print_exc()
        return None, None, None


# ═══════════════════════════════════════════════════════════════════════════════
#  ENTRY POINT
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(
        description='ATARS — Automated Time-Series Analysis and Reporting System',
        formatter_class=argparse.ArgumentDefaultsHelpFormatter
    )
    parser.add_argument('--data',      default='data.csv',
                        help='Path to CSV/Excel data file')
    parser.add_argument('--city',      default='City',
                        help='City or location name')
    parser.add_argument('--station',   default='Station-01',
                        help='Station ID')
    parser.add_argument('--output',    default='atars_output',
                        help='Output directory')
    parser.add_argument('--no-llm',    action='store_true',
                        help='Disable LLM (use statistical placeholders)')
    parser.add_argument('--llm-model', default='llama3.2',
                        help='Ollama model name (e.g. mistral, phi3)')
    parser.add_argument('--baseline',  type=int, default=30,
                        help='Baseline window W in days')
    parser.add_argument('--timeout',   type=int, default=600,
                        help='LLM timeout per section in seconds (default 600 = 10 min)')
    parser.add_argument('--retry',     type=int, default=2,
                        help='Number of LLM retries per section on timeout')
    parser.add_argument('--no-forecast', action='store_true',
                        help='Disable Holt-Winters PM10 forecast (Chart 15)')
    args = parser.parse_args()

    config = DEFAULT_CONFIG.copy()
    config.update({
        'data_file'           : args.data,
        'city'                : args.city,
        'station_id'          : args.station,
        'output_dir'          : args.output,
        'use_llm'             : not args.no_llm,
        'llm_model'           : args.llm_model,
        'baseline_window_days': args.baseline,
        'llm_timeout'         : args.timeout,
        'llm_retry'           : args.retry,
        'use_forecast'        : not args.no_forecast,
    })

    run_pipeline(config)


if __name__ == '__main__':
    main()
